import streamlit as st
import os
import html
import sys  # for adding summarizer path
import subprocess
import pickle
import random
import re
import datetime
from datetime import datetime, date, timedelta
from PIL import Image
import calendar
import time
import json
import hashlib
import base64
import io
from io import BytesIO
import spacy
import spacy.cli
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
from PyPDF2 import PdfReader
from docx import Document
import speech_recognition as sr
import tempfile
from pydub import AudioSegment
import pyshorteners
import pyperclip as clipboard
import fitz  # PyMuPDF for PDF image extraction
import zipfile
import xml.etree.ElementTree as ET
import streamlit.components.v1 as components
import pyttsx3
import threading
import queue




# Add the Text-Summarizer folder to sys.path so we can import summarizer
# summ_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'Text-Summarizer'))
# if summ_path not in sys.path:
#     sys.path.insert(0, summ_path)






# === Utility functions for PDF/DOCX text extraction ===
# === Text Summarizer (Integrated) ===
# Download spacy model if not already downloaded
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    spacy.cli.download("en_core_web_sm")
    nlp = spacy.load('en_core_web_sm')






# === Voice activation //////////////////////////////////////////////////////////////////////////////////////===
class VoiceActivation:
    def __init__(self):
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        self.tts_engine = pyttsx3.init()
        self.is_listening = False
        self.voice_queue = queue.Queue()
        self.setup_tts()
        
    def setup_tts(self):
        """Setup text-to-speech engine"""
        try:
            # Set voice properties
            voices = self.tts_engine.getProperty('voices')
            if voices:
                # Try to set a female voice if available
                for voice in voices:
                    if 'female' in voice.name.lower() or 'zira' in voice.name.lower():
                        self.tts_engine.setProperty('voice', voice.id)
                        break
                else:
                    # Use first available voice
                    self.tts_engine.setProperty('voice', voices[0].id)
            
            # Set speech rate and volume
            self.tts_engine.setProperty('rate', 180)  # Speed of speech
            self.tts_engine.setProperty('volume', 0.9)  # Volume level (0.0 to 1.0)
        except Exception as e:
            st.error(f"Error setting up text-to-speech: {e}")
    
    def speak(self, text):
        """Convert text to speech"""
        try:
            self.tts_engine.say(text)
            self.tts_engine.runAndWait()
        except Exception as e:
            st.error(f"Error in text-to-speech: {e}")
    
    def listen_for_wake_word(self, app_name):
        """Listen for the wake word (app name)"""
        try:
            with self.microphone as source:
                # Adjust for ambient noise
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
            
            while self.is_listening:
                try:
                    with self.microphone as source:
                        # Listen for audio with timeout
                        audio = self.recognizer.listen(source, timeout=1, phrase_time_limit=3)
                    
                    # Recognize speech
                    text = self.recognizer.recognize_google(audio).lower()
                    
                    # Check for wake word
                    if f"hey {app_name.lower()}" in text:
                        self.voice_queue.put(("wake_word_detected", text))
                        
                        # Listen for command after wake word
                        time.sleep(0.5)  # Brief pause
                        with self.microphone as source:
                            audio = self.recognizer.listen(source, timeout=5, phrase_time_limit=5)
                        
                        command = self.recognizer.recognize_google(audio).lower()
                        self.voice_queue.put(("command", command))
                
                except sr.WaitTimeoutError:
                    # Timeout is normal, continue listening
                    continue
                except sr.UnknownValueError:
                    # Could not understand audio, continue listening
                    continue
                except sr.RequestError as e:
                    self.voice_queue.put(("error", f"Speech recognition error: {e}"))
                    break
                except Exception as e:
                    self.voice_queue.put(("error", f"Unexpected error: {e}"))
                    break
        
        except Exception as e:
            self.voice_queue.put(("error", f"Microphone error: {e}"))
    
    def start_listening(self, app_name):
        """Start listening for voice commands"""
        if not self.is_listening:
            self.is_listening = True
            # Start listening in a separate thread
            listening_thread = threading.Thread(
                target=self.listen_for_wake_word, 
                args=(app_name,),
                daemon=True
            )
            listening_thread.start()
    
    def stop_listening(self):
        """Stop listening for voice commands"""
        self.is_listening = False
    
    def process_command(self, command, user_name):
        """Process voice command and return appropriate response and action"""
        command = command.lower().strip()
        
        # Define command mappings
        command_mappings = {
            # Navigation commands
            'home': ('📅 Home', f"Taking you to Home, {user_name}"),
            'calendar': ('📅 Home', f"Opening Calendar, {user_name}"),
            'dashboard': ('📅 Home', f"Opening Dashboard, {user_name}"),
            
            # Text Summarizer
            'text summarizer': ('📝 Text Summarizer', f"Activating Text Summarizer, {user_name}"),
            'summarizer': ('📝 Text Summarizer', f"Opening Text Summarizer, {user_name}"),
            'summarize': ('📝 Text Summarizer', f"Opening Text Summarizer, {user_name}"),
            'summary': ('📝 Text Summarizer', f"Opening Text Summarizer, {user_name}"),
            
            # Flashcards
            'flashcards': ('🃏 Flashcard Generator', f"Activating Flashcards, {user_name}"),
            'flashcard': ('🃏 Flashcard Generator', f"Opening Flashcard Generator, {user_name}"),
            'cards': ('🃏 Flashcard Generator', f"Opening Flashcards, {user_name}"),
            'flash cards': ('🃏 Flashcard Generator', f"Activating Flashcards, {user_name}"),
            
            # MCQ Test
            'mcq': ('🧠 MCQ Test Generator', f"Opening MCQ Test Generator, {user_name}"),
            'test': ('🧠 MCQ Test Generator', f"Activating Test Generator, {user_name}"),
            'quiz': ('🧠 MCQ Test Generator', f"Opening Quiz Generator, {user_name}"),
            'mcq test': ('🧠 MCQ Test Generator', f"Activating MCQ Test Generator, {user_name}"),
            
            # URL Shortener
            'url shortener': ('🔗 URL Shortener', f"Opening URL Shortener, {user_name}"),
            'shortener': ('🔗 URL Shortener', f"Activating URL Shortener, {user_name}"),
            'short url': ('🔗 URL Shortener', f"Opening URL Shortener, {user_name}"),
            'link shortener': ('🔗 URL Shortener', f"Opening Link Shortener, {user_name}"),
            
            # Grade Calculator
            'grade calculator': ('📊 Grade Calculator', f"Opening Grade Calculator, {user_name}"),
            'grades': ('📊 Grade Calculator', f"Activating Grade Calculator, {user_name}"),
            'calculate grade': ('📊 Grade Calculator', f"Opening Grade Calculator, {user_name}"),
            'grade': ('📊 Grade Calculator', f"Opening Grade Calculator, {user_name}"),
            
            # SGPA Calculator
            'sgpa': ('🎓 SGPA Calculator', f"Opening SGPA Calculator, {user_name}"),
            'sgpa calculator': ('🎓 SGPA Calculator', f"Activating SGPA Calculator, {user_name}"),
            'semester gpa': ('🎓 SGPA Calculator', f"Opening SGPA Calculator, {user_name}"),
            
            # CGPA Calculator
            'cgpa': ('CGPA Calculator', f"Opening CGPA Calculator, {user_name}"),
            'cgpa calculator': ('CGPA Calculator', f"Activating CGPA Calculator, {user_name}"),
            'cumulative gpa': ('CGPA Calculator', f"Opening CGPA Calculator, {user_name}"),
            
            # Physics Solver
            'physics': ('⚡ Physics Solver', f"Opening Physics Solver, {user_name}"),
            'physics solver': ('⚡ Physics Solver', f"Activating Physics Solver, {user_name}"),
            'solve physics': ('⚡ Physics Solver', f"Opening Physics Solver, {user_name}"),
            
            # Unit Converter
            'unit converter': ('📏 Unit Converter', f"Opening Unit Converter, {user_name}"),
            'converter': ('📏 Unit Converter', f"Activating Unit Converter, {user_name}"),
            'convert units': ('📏 Unit Converter', f"Opening Unit Converter, {user_name}"),
            'units': ('📏 Unit Converter', f"Opening Unit Converter, {user_name}"),
        }
        
        # Find matching command
        for key, (page, response) in command_mappings.items():
            if key in command:
                return page, response
        
        # If no specific command found, provide help
        return None, f"I didn't understand that command, {user_name}. Try saying things like 'open flashcards' or 'activate summarizer'."

def voice_settings_page():
    """Voice activation settings page"""
    st.header("🎤 Voice Activation Settings")
    
    # Initialize session state for voice settings
    if 'voice_activated' not in st.session_state:
        st.session_state.voice_activated = False
    if 'app_name' not in st.session_state:
        st.session_state.app_name = ""
    if 'user_name' not in st.session_state:
        st.session_state.user_name = ""
    if 'voice_assistant' not in st.session_state:
        st.session_state.voice_assistant = None
    
    # Voice activation toggle
    st.subheader("🔊 Voice Control")
    
    voice_enabled = st.toggle(
        "Enable Voice Activation", 
        value=st.session_state.voice_activated,
        help="Turn on voice commands to control the app with your voice"
    )
    
    # Update session state
    if voice_enabled != st.session_state.voice_activated:
        st.session_state.voice_activated = voice_enabled
        
        if not voice_enabled and st.session_state.voice_assistant:
            # Stop voice assistant if disabled
            st.session_state.voice_assistant.stop_listening()
            st.session_state.voice_assistant = None
    
    if voice_enabled:
        st.success("🎤 Voice activation is ON")
        
        # Required fields when voice is enabled
        st.subheader("📝 Required Settings")
        
        col1, col2 = st.columns(2)
        
        with col1:
            app_name = st.text_input(
                "App Name *", 
                value=st.session_state.app_name,
                placeholder="e.g., Buddy",
                help="This is the wake word. Say 'Hey [App Name]' to activate voice commands"
            )
        
        with col2:
            user_name = st.text_input(
                "Your Name *", 
                value=st.session_state.user_name,
                placeholder="e.g., Uma",
                help="The app will address you by this name"
            )
        
        # Validation
        if not app_name.strip() or not user_name.strip():
            st.error("⚠️ Both App Name and Your Name are required when voice activation is enabled!")
            st.session_state.voice_activated = False
            return
        
        # Update session state
        st.session_state.app_name = app_name.strip()
        st.session_state.user_name = user_name.strip()
        
        # Voice assistant setup
        if st.session_state.voice_activated and not st.session_state.voice_assistant:
            try:
                st.session_state.voice_assistant = VoiceActivation()
                st.session_state.voice_assistant.start_listening(st.session_state.app_name)
                st.success(f"🎤 Voice assistant started! Say 'Hey {st.session_state.app_name}' to begin.")
            except Exception as e:
                st.error(f"Failed to start voice assistant: {e}")
                st.session_state.voice_activated = False
        
        # Instructions
        st.subheader("📋 How to Use Voice Commands")
        
        st.markdown(f"""
        **Wake Word:** Say "Hey {app_name}" to activate voice commands
        
        **Example Commands:**
        - "Hey {app_name}" → App responds: "Hey {user_name}"
        - Then say: "Activate flashcards" → App responds: "Activating Flashcards" and opens the feature
        
        **Available Commands:**
        - **Home/Calendar:** "home", "calendar", "dashboard"
        - **Text Summarizer:** "text summarizer", "summarizer", "summarize"
        - **Flashcards:** "flashcards", "flashcard", "cards"
        - **MCQ Test:** "mcq", "test", "quiz"
        - **URL Shortener:** "url shortener", "shortener"
        - **Grade Calculator:** "grade calculator", "grades"
        - **SGPA Calculator:** "sgpa", "sgpa calculator"
        - **CGPA Calculator:** "cgpa", "cgpa calculator"
        - **Physics Solver:** "physics", "physics solver"
        - **Unit Converter:** "unit converter", "converter"
        """)
        
        # Test voice feature
        st.subheader("🧪 Test Voice Feature")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🎤 Test Voice Response"):
                if st.session_state.voice_assistant:
                    try:
                        st.session_state.voice_assistant.speak(f"Hey {user_name}! Voice activation is working perfectly!")
                        st.success("✅ Voice test completed!")
                    except Exception as e:
                        st.error(f"Voice test failed: {e}")
                else:
                    st.error("Voice assistant not initialized")
        
        with col2:
            if st.button("🔄 Restart Voice Assistant"):
                try:
                    if st.session_state.voice_assistant:
                        st.session_state.voice_assistant.stop_listening()
                    
                    st.session_state.voice_assistant = VoiceActivation()
                    st.session_state.voice_assistant.start_listening(st.session_state.app_name)
                    st.success("🎤 Voice assistant restarted!")
                except Exception as e:
                    st.error(f"Failed to restart voice assistant: {e}")
        
        # Voice status indicator
        if st.session_state.voice_assistant and st.session_state.voice_assistant.is_listening:
            st.info(f"🎤 Listening for 'Hey {app_name}'...")
        
    else:
        st.info("🔇 Voice activation is OFF")
        
        # Stop voice assistant if it's running
        if st.session_state.voice_assistant:
            st.session_state.voice_assistant.stop_listening()
            st.session_state.voice_assistant = None
    
    # Additional settings
    st.subheader("⚙️ Advanced Settings")
    
    with st.expander("🔧 Voice Recognition Settings"):
        st.markdown("""
        **Microphone Requirements:**
        - Ensure your microphone is connected and working
        - Grant microphone permissions to your browser/app
        - Speak clearly and at a normal pace
        - Minimize background noise for better recognition
        
        **Troubleshooting:**
        - If voice commands aren't working, try restarting the voice assistant
        - Check your internet connection (required for speech recognition)
        - Make sure your microphone isn't muted
        """)

def handle_voice_commands():
    """Handle voice commands in the main app"""
    if (st.session_state.get('voice_activated', False) and 
        st.session_state.get('voice_assistant') and 
        st.session_state.voice_assistant.is_listening):
        
        # Check for voice commands
        try:
            while not st.session_state.voice_assistant.voice_queue.empty():
                event_type, data = st.session_state.voice_assistant.voice_queue.get_nowait()
                
                if event_type == "wake_word_detected":
                    # Respond to wake word
                    user_name = st.session_state.get('user_name', 'User')
                    st.session_state.voice_assistant.speak(f"Hey {user_name}")
                    st.success(f"🎤 Wake word detected: {data}")
                
                elif event_type == "command":
                    # Process command
                    user_name = st.session_state.get('user_name', 'User')
                    page, response = st.session_state.voice_assistant.process_command(data, user_name)
                    
                    # Speak response
                    st.session_state.voice_assistant.speak(response)
                    
                    if page:
                        # Navigate to the requested page
                        st.session_state.current_page = page
                        st.success(f"🎤 Command executed: {data} → {page}")
                        st.rerun()
                    else:
                        st.warning(f"🎤 Command not recognized: {data}")
                
                elif event_type == "error":
                    st.error(f"🎤 Voice error: {data}")
        
        except queue.Empty:
            pass
        except Exception as e:
            st.error(f"Error handling voice commands: {e}")

def add_voice_to_sidebar():
    """Add voice activation indicator to sidebar"""
    if st.session_state.get('voice_activated', False):
        app_name = st.session_state.get('app_name', 'App')
        user_name = st.session_state.get('user_name', 'User')
        
        st.sidebar.markdown("---")
        st.sidebar.markdown("### 🎤 Voice Control")
        
        if (st.session_state.get('voice_assistant') and 
            st.session_state.voice_assistant.is_listening):
            st.sidebar.success(f"🟢 Listening for 'Hey {app_name}'")
        else:
            st.sidebar.error("🔴 Voice assistant offline")
        
        st.sidebar.caption(f"Configured for: {user_name}")

# Integration function to add to your main app
def integrate_voice_activation():
    """
    Call this function in your main app to integrate voice activation
    Add this to your main() function after the sidebar setup
    """
    # Handle voice commands
    handle_voice_commands()
    
    # Add voice indicator to sidebar
    add_voice_to_sidebar()




# === Text Summarizer Function //////////////////////////////////////////////////////////////////////////////////////===
def summarizer(rawdocs):
    """
    Summarizes the input text using spaCy NLP
    Returns: summary, original_doc, original_word_count, summary_word_count
    """
    stopwords = list(STOP_WORDS)
    
    doc = nlp(rawdocs)
    
    # Calculate word frequencies
    word_freq = {}
    for word in doc:
        if word.text.lower() not in stopwords and word.text.lower() not in punctuation:
            if word.text not in word_freq.keys():
                word_freq[word.text] = 1
            else:
                word_freq[word.text] += 1
    
    # Normalize frequencies
    if word_freq:  # Check if word_freq is not empty
        max_freq = max(word_freq.values())
        for word in word_freq.keys():
            word_freq[word] = word_freq[word] / max_freq
    
    # Calculate sentence scores
    sent_tokens = [sent for sent in doc.sents]
    sent_scores = {}
    
    for sent in sent_tokens:
        for word in sent:
            if word.text in word_freq.keys():
                if sent not in sent_scores.keys():
                    sent_scores[sent] = word_freq[word.text]
                else:
                    sent_scores[sent] += word_freq[word.text]
    
    # Select top 30% sentences for summary
    if sent_scores:
        select_len = max(1, int(len(sent_tokens) * 0.3))  # Ensure at least 1 sentence
        summary_sents = nlargest(select_len, sent_scores, key=sent_scores.get)
        
        # Sort sentences by their original order in the text
        summary_sents_ordered = sorted(summary_sents, key=lambda x: list(doc.sents).index(x))
        final_summary = ' '.join([sent.text for sent in summary_sents_ordered])
    else:
        final_summary = rawdocs[:200] + "..." if len(rawdocs) > 200 else rawdocs
    
    return final_summary, rawdocs, len(rawdocs.split(' ')), len(final_summary.split(' '))

# === File Text Extraction Functions ===
def extract_text_from_pdf(uploaded_file):
    """Extract text from PDF file"""
    try:
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(uploaded_file):
    """Extract text from DOCX file"""
    try:
        doc = Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

# === Voice Recognition Function ===
def transcribe_audio(audio_file):
    """
    Transcribe audio file to text using speech recognition
    """
    try:
        recognizer = sr.Recognizer()
        
        # Create a temporary file to save the uploaded audio
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            # If the uploaded file is not WAV, convert it
            if audio_file.type != "audio/wav":
                # Read the uploaded file
                audio_data = audio_file.read()
                
                # Convert to AudioSegment
                audio = AudioSegment.from_file(io.BytesIO(audio_data))
                
                # Export as WAV
                audio.export(tmp_file.name, format="wav")
            else:
                # If it's already WAV, just write it
                tmp_file.write(audio_file.read())
            
            tmp_file_path = tmp_file.name
        
        # Transcribe the audio
        with sr.AudioFile(tmp_file_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        
        return text
    
    except sr.UnknownValueError:
        return "Could not understand the audio. Please try again with clearer audio."
    except sr.RequestError as e:
        return f"Error with speech recognition service: {str(e)}"
    except Exception as e:
        return f"Error processing audio: {str(e)}"

# === Main Text Summarizer Page ===
def text_summarizer_page():
    st.header("📝 Text Summarizer")
    st.markdown("Upload files, paste text, or record audio to get instant summaries!")
    
    # Create tabs for different input methods
    tab1, tab2, tab3 = st.tabs(["📄 File Upload", "✍️ Text Input", "🎤 Voice Input"])
    
    text_to_summarize = ""
    
    with tab1:
        st.subheader("Upload Document")
        uploaded_file = st.file_uploader(
            "Choose a file", 
            type=['pdf', 'docx', 'txt'],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        if uploaded_file is not None:
            if uploaded_file.type == "application/pdf":
                text_to_summarize = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                text_to_summarize = extract_text_from_docx(uploaded_file)
            elif uploaded_file.type == "text/plain":
                text_to_summarize = uploaded_file.getvalue().decode("utf-8")
            
            if text_to_summarize:
                st.success(f"✅ File loaded successfully! ({len(text_to_summarize.split())} words)")
                with st.expander("Preview extracted text"):
                    st.text_area("Extracted Text", text_to_summarize[:1000] + "..." if len(text_to_summarize) > 1000 else text_to_summarize, height=200)
    
    with tab2:
        st.subheader("Paste Text")
        text_input = st.text_area(
            "Enter your text here:", 
            height=200,
            placeholder="Paste your text here for summarization..."
        )
        if text_input.strip():
            text_to_summarize = text_input
            st.success(f"✅ Text entered! ({len(text_to_summarize.split())} words)")
    
    with tab3:
        st.subheader("Voice Recording")
        st.info("Upload an audio file to transcribe and summarize")
        
        audio_file = st.file_uploader(
            "Choose an audio file",
            type=['wav', 'mp3', 'mp4', 'm4a', 'ogg'],
            help="Supported formats: WAV, MP3, MP4, M4A, OGG"
        )
        
        if audio_file is not None:
            st.audio(audio_file, format='audio/wav')
            
            if st.button("🎤 Transcribe Audio", key="transcribe_btn"):
                with st.spinner("Transcribing audio... This may take a moment."):
                    transcribed_text = transcribe_audio(audio_file)
                    if transcribed_text and not transcribed_text.startswith("Could not") and not transcribed_text.startswith("Error"):
                        text_to_summarize = transcribed_text
                        st.success("✅ Audio transcribed successfully!")
                        st.subheader("Transcribed Text:")
                        st.text_area("", transcribed_text, height=150)
                    else:
                        st.error(transcribed_text)
    
    # Summarization section
    if text_to_summarize:
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("✨ Generate Summary", type="primary", use_container_width=True):
                if not text_to_summarize.strip():
                    st.error("❌ Please provide some text to summarize.")
                else:
                    with st.spinner("Generating summary..."):
                        summary, original_text, original_count, summary_count = summarizer(text_to_summarize)
                        
                        # Display results in two columns
                        st.markdown("## 📊 Results")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.subheader("📄 Original Text")
                            st.text_area(
                                "", 
                                original_text, 
                                height=400,
                                key="original_text"
                            )
                            st.metric("Word Count", original_count)
                        
                        with col2:
                            st.subheader("📝 Summary")
                            st.text_area(
                                "", 
                                summary, 
                                height=400,
                                key="summary_text"
                            )
                            st.metric("Word Count", summary_count)
                        
                        # Summary statistics
                        reduction_percentage = round((1 - summary_count/original_count) * 100, 1) if original_count > 0 else 0
                        
                        st.markdown("### 📈 Summary Statistics")
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Original Words", original_count)
                        with col2:
                            st.metric("Summary Words", summary_count)
                        with col3:
                            st.metric("Reduction", f"{reduction_percentage}%")
                        
                        # Download buttons
                        st.markdown("### 💾 Download Options")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                label="📄 Download Original Text",
                                data=original_text,
                                file_name="original_text.txt",
                                mime="text/plain"
                            )
                        with col2:
                            st.download_button(
                                label="📝 Download Summary",
                                data=summary,
                                file_name="summary.txt",
                                mime="text/plain"
                            )
    else:
        st.info("👆 Please upload a file, enter text, or upload an audio recording to get started!")

# === Installation Instructions ===
def show_installation_requirements():
    st.sidebar.markdown("---")
    with st.sidebar.expander("📦 Installation Requirements"):
        st.markdown("""
        To run this app, install the required packages:
        
        ```bash
        pip install streamlit
        pip install spacy
        pip install PyPDF2
        pip install python-docx
        pip install SpeechRecognition
        pip install pydub
        python -m spacy download en_core_web_sm
        ```
        
        For audio processing, you may also need:
        ```bash
        # For various audio format support
        pip install pydub[mp3]
        
        # For speech recognition
        pip install pyaudio  # (optional, for microphone input)
        ```
        """)




# === Task tracker //////////////////////////////////////////////////////////////////////////////////////===
# === Data persistence functions ===
def save_todos_to_file(todos):
    """Save todos to a pickle file for persistence"""
    try:
        with open('study_buddy_todos.pkl', 'wb') as f:
            pickle.dump(todos, f)
    except Exception as e:
        st.error(f"Error saving todos: {e}")

def load_todos_from_file():
    """Load todos from pickle file"""
    try:
        if os.path.exists('study_buddy_todos.pkl'):
            with open('study_buddy_todos.pkl', 'rb') as f:
                return pickle.load(f)
    except Exception as e:
        st.error(f"Error loading todos: {e}")
    return []


# Set page config for a wider layout and custom title
st.set_page_config(
    page_title="Study Buddy 📚",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced CSS for color-coded calendar and task management
st.markdown("""
<style>
.stButton button {
    background-color: #FF6B6B;
    color: white;
    border-radius: 20px;
    padding: 0.5rem 1rem;
}
            
input {
    color: black !important;
}

.stButton button:hover {
    background-color: #FF8787;
}
.stTextInput input, .stTextArea textarea, .stSelectbox select {
    border-radius: 15px;
}
.task-done {
    text-decoration: line-through;
    color: #888888;
}
.calendar-day {
    border: 1px solid #ddd;
    border-radius: 8px;
    padding: 8px;
    margin: 2px;
    min-height: 120px;
    background-color: #f8f9fa;
    cursor: pointer;
    transition: all 0.3s ease;
    position: relative;
}
.calendar-day:hover {
    transform: scale(1.02);
    box-shadow: 0 4px 8px rgba(0,0,0,0.1);
}
.calendar-day.today {
    background-color: #e3f2fd;
    border-color: #2196f3;
}
.calendar-day.has-tasks {
    border-width: 2px;
}
/* Color coding based on deadline urgency */
.calendar-day.urgent {
    background-color: #ffebee;
    border-color: #f44336;
}
.calendar-day.warning {
    background-color: #fff8e1;
    border-color: #ff9800;
}
.calendar-day.safe {
    background-color: #e8f5e8;
    border-color: #4caf50;
}
.day-number {
    font-size: 18px;
    font-weight: bold;
    color: #333;
    margin-bottom: 8px;
    display: block;
}
.day-content {
    margin-top: 4px;
}
.task-item {
    font-size: 0.8em;
    padding: 2px 4px;
    margin: 1px 0;
    border-radius: 4px;
    background-color: #f5f5f5;
    cursor: pointer;
}
.task-high {
    background-color: #ffebee;
    border-left: 3px solid #f44336;
}
.task-medium {
    background-color: #fff8e1;
    border-left: 3px solid #ff9800;
}
.task-low {
    background-color: #e8f5e8;
    border-left: 3px solid #4caf50;
}
.task-popup {
    position: fixed;
    top: 50%;
    left: 50%;
    transform: translate(-50%, -50%);
    background: white;
    border: 2px solid #ddd;
    border-radius: 15px;
    padding: 20px;
    box-shadow: 0 8px 16px rgba(0,0,0,0.2);
    z-index: 1000;
    max-width: 500px;
    width: 90%;
}
.overlay {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    background: rgba(0,0,0,0.5);
    z-index: 999;
}
</style>
""", unsafe_allow_html=True)

# === Todo List with persistence ===
def load_todos():
    if 'todos' not in st.session_state:
        st.session_state.todos = load_todos_from_file()
    return st.session_state.todos

def save_todos(todos):
    st.session_state.todos = todos
    save_todos_to_file(todos)

def get_days_until_deadline(due_date_str):
    """Calculate days until deadline"""
    try:
        due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
        today = date.today()
        return (due_date - today).days
    except:
        return 999

def get_urgency_class(days_until):
    """Return CSS class based on urgency"""
    if days_until < 0:
        return "urgent"  # Overdue
    elif 0 <= days_until <= 3:
        return "urgent"  # 0-3 days (red)
    elif 4 <= days_until <= 6:
        return "warning"  # 4-6 days (yellow)
    else:
        return "safe"  # 7+ days (green)

# === Calendar Home Page with Enhanced Features ===
def calendar_home():
    st.header("📅 Your Study Calendar")
    
    # Load todos
    todos = load_todos()
    
    # Initialize session state for popup management
    if 'show_task_popup' not in st.session_state:
        st.session_state.show_task_popup = False
    if 'selected_date' not in st.session_state:
        st.session_state.selected_date = None
    if 'popup_tasks' not in st.session_state:
        st.session_state.popup_tasks = []
    
    # Get current date and month navigation
    if 'current_month' not in st.session_state:
        st.session_state.current_month = datetime.now().replace(day=1)
    
    col1, col2, col3, col4, col5 = st.columns([1, 1, 2, 1, 1])
    
    with col1:
        if st.button("◀ Previous"):
            if st.session_state.current_month.month == 1:
                st.session_state.current_month = st.session_state.current_month.replace(year=st.session_state.current_month.year - 1, month=12)
            else:
                st.session_state.current_month = st.session_state.current_month.replace(month=st.session_state.current_month.month - 1)
            st.rerun()
    
    with col2:
        if st.button("Today"):
            st.session_state.current_month = datetime.now().replace(day=1)
            st.rerun()
    
    with col3:
        st.markdown(f"### {st.session_state.current_month.strftime('%B %Y')}")
    
    with col4:
        if st.button("Next ▶"):
            if st.session_state.current_month.month == 12:
                st.session_state.current_month = st.session_state.current_month.replace(year=st.session_state.current_month.year + 1, month=1)
            else:
                st.session_state.current_month = st.session_state.current_month.replace(month=st.session_state.current_month.month + 1)
            st.rerun()
    
    with col5:
        if st.button("➕ Add Task"):
            st.session_state.show_add_task = True
            st.rerun()
    
    # Quick add task form
    if getattr(st.session_state, 'show_add_task', False):
        with st.expander("Add New Task", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                new_task = st.text_input("Task Description", key="quick_task")
                category = st.selectbox("Category", ["📚 Study", "📝 Assignment", "📖 Reading", "🎯 Project", "🎨 Other"], key="quick_category")
            
            with col2:
                priority = st.selectbox("Priority", ["🔥 High", "⚡ Medium", "💫 Low"], key="quick_priority")
                due_date = st.date_input("Due Date", min_value=date.today(), key="quick_due_date")
            
            col_add, col_cancel = st.columns(2)
            with col_add:
                if st.button("Add Task", key="confirm_add"):
                    if new_task:
                        todos.append({
                            "id": len(todos) + 1,  # Simple ID system
                            "task": new_task,
                            "category": category,
                            "priority": priority,
                            "due_date": due_date.strftime("%Y-%m-%d"),
                            "completed": False,
                            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        })
                        save_todos(todos)
                        st.session_state.show_add_task = False
                        st.success("Task added successfully!")
                        st.rerun()
            
            with col_cancel:
                if st.button("Cancel", key="cancel_add"):
                    st.session_state.show_add_task = False
                    st.rerun()
    
    # Create calendar grid
    current_date = datetime.now().date()
    cal = calendar.monthcalendar(st.session_state.current_month.year, st.session_state.current_month.month)
    
    # Days of week header
    days_of_week = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    cols = st.columns(7)
    for i, day in enumerate(days_of_week):
        with cols[i]:
            st.markdown(f"**{day}**")
    
    # Calendar days with color coding
    for week in cal:
        cols = st.columns(7)
        for i, day in enumerate(week):
            with cols[i]:
                if day == 0:
                    st.markdown("<div style='height: 120px;'></div>", unsafe_allow_html=True)
                else:
                    # Check if this day has tasks
                    day_date = date(st.session_state.current_month.year, st.session_state.current_month.month, day)
                    day_tasks = [todo for todo in todos if todo['due_date'] == day_date.strftime("%Y-%m-%d") and not todo['completed']]
                    
                    # Determine day style with color coding
                    day_class = "calendar-day"
                    if day_date == current_date:
                        day_class += " today"
                    
                    if day_tasks:
                        day_class += " has-tasks"
                        # Get the most urgent task for color coding
                        min_days = min([get_days_until_deadline(task['due_date']) for task in day_tasks])
                        urgency_class = get_urgency_class(min_days)
                        day_class += f" {urgency_class}"
                    
                    # Create day container with better layout
                    st.markdown(f"""
                    <div class='{day_class}'>
                        <div class='day-number'>{day}</div>
                        <div class='day-content'>
                    """, unsafe_allow_html=True)
                    
                    # Add task previews
                    for task in day_tasks[:2]:  # Show max 2 tasks for preview
                        priority_class = "task-low"
                        if task['priority'] == "🔥 High":
                            priority_class = "task-high"
                        elif task['priority'] == "⚡ Medium":
                            priority_class = "task-medium"
                        
                        task_text = task['task']
                        if len(task_text) > 12:
                            task_text = task_text[:12] + "..."
                        
                        st.markdown(f"<div class='task-item {priority_class}'>{task_text}</div>", 
                                  unsafe_allow_html=True)
                    
                    if len(day_tasks) > 2:
                        st.markdown(f"<div class='task-item'>+{len(day_tasks) - 2} more</div>", 
                                  unsafe_allow_html=True)
                    
                    st.markdown("</div></div>", unsafe_allow_html=True)
                    
                    # Clickable button for day management
                    day_key = f"day_btn_{day_date.strftime('%Y-%m-%d')}"
                    if st.button("📅 Manage", key=day_key, help=f"Click to manage tasks for {day_date.strftime('%B %d, %Y')}"):
                        st.session_state.selected_date = day_date
                        st.session_state.popup_tasks = [todo for todo in todos if todo['due_date'] == day_date.strftime("%Y-%m-%d")]
                        st.session_state.show_task_popup = True
                        st.rerun()
    
    # Task Management Popup
    if st.session_state.show_task_popup and st.session_state.selected_date:
        selected_date = st.session_state.selected_date
        popup_tasks = st.session_state.popup_tasks
        
        st.markdown("---")
        st.markdown(f"### 📋 Tasks for {selected_date.strftime('%B %d, %Y')}")
        
        if popup_tasks:
            for task_idx, task in enumerate(popup_tasks):
                with st.container():
                    col1, col2, col3, col4, col5 = st.columns([3, 1, 1, 1, 1])
                    
                    with col1:
                        status = "✅ Completed" if task['completed'] else "⏳ Schedule Pending"
                        st.write(f"**{task['task']}** - {status}")
                        st.write(f"{task['category']} | {task['priority']}")
                    
                    with col2:
                        if st.button("✅ Complete", key=f"complete_{task['id']}"):
                            # Toggle completion status
                            for i, todo in enumerate(todos):
                                if todo['id'] == task['id']:
                                    todos[i]['completed'] = not todos[i]['completed']
                                    break
                            save_todos(todos)
                            st.success("Task status updated!")
                            st.rerun()
                    
                    with col3:
                        if st.button("🗑️ Delete", key=f"delete_{task['id']}"):
                            todos = [todo for todo in todos if todo['id'] != task['id']]
                            save_todos(todos)
                            st.success("Task deleted!")
                            st.rerun()
                    
                    with col4:
                        if st.button("📅 Change Date", key=f"change_date_{task['id']}"):
                            st.session_state[f'editing_date_{task["id"]}'] = True
                            st.rerun()
                    
                    with col5:
                        if st.button("✏️ Edit", key=f"edit_{task['id']}"):
                            st.session_state[f'editing_task_{task["id"]}'] = True
                            st.rerun()
                    
                    # Date change form
                    if st.session_state.get(f'editing_date_{task["id"]}', False):
                        with st.form(f"change_date_form_{task['id']}"):
                            new_date = st.date_input("New Due Date", 
                                                   value=datetime.strptime(task['due_date'], "%Y-%m-%d").date(),
                                                   min_value=date.today())
                            
                            col_save, col_cancel = st.columns(2)
                            with col_save:
                                submitted = st.form_submit_button("Save Date")
                                if submitted:
                                    for i, todo in enumerate(todos):
                                        if todo['id'] == task['id']:
                                            todos[i]['due_date'] = new_date.strftime("%Y-%m-%d")
                                            break
                                    save_todos(todos)
                                    st.session_state[f'editing_date_{task["id"]}'] = False
                                    st.success("Date updated!")
                                    st.rerun()
                            
                            with col_cancel:
                                if st.form_submit_button("Cancel"):
                                    st.session_state[f'editing_date_{task["id"]}'] = False
                                    st.rerun()
                    
                    # Task edit form
                    if st.session_state.get(f'editing_task_{task["id"]}', False):
                        with st.form(f"edit_task_form_{task['id']}"):
                            new_task_text = st.text_input("Task Description", value=task['task'])
                            new_category = st.selectbox("Category", 
                                                      ["📚 Study", "📝 Assignment", "📖 Reading", "🎯 Project", "🎨 Other"],
                                                      index=["📚 Study", "📝 Assignment", "📖 Reading", "🎯 Project", "🎨 Other"].index(task['category']))
                            new_priority = st.selectbox("Priority", 
                                                      ["🔥 High", "⚡ Medium", "💫 Low"],
                                                      index=["🔥 High", "⚡ Medium", "💫 Low"].index(task['priority']))
                            
                            col_save, col_cancel = st.columns(2)
                            with col_save:
                                submitted = st.form_submit_button("Save Changes")
                                if submitted:
                                    for i, todo in enumerate(todos):
                                        if todo['id'] == task['id']:
                                            todos[i]['task'] = new_task_text
                                            todos[i]['category'] = new_category
                                            todos[i]['priority'] = new_priority
                                            break
                                    save_todos(todos)
                                    st.session_state[f'editing_task_{task["id"]}'] = False
                                    st.success("Task updated!")
                                    st.rerun()
                            
                            with col_cancel:
                                if st.form_submit_button("Cancel"):
                                    st.session_state[f'editing_task_{task["id"]}'] = False
                                    st.rerun()
                    
                    st.markdown("---")
        else:
            st.info("No tasks scheduled for this date.")
            
            # Option to add task for selected date
            if st.button("➕ Add Task for This Date"):
                st.session_state.quick_add_for_date = selected_date
                st.rerun()
        
        # Quick add for selected date
        if st.session_state.get('quick_add_for_date'):
            with st.form("quick_add_form"):
                st.subheader(f"Add Task for {selected_date.strftime('%B %d, %Y')}")
                
                task_desc = st.text_input("Task Description")
                category = st.selectbox("Category", ["📚 Study", "📝 Assignment", "📖 Reading", "🎯 Project", "🎨 Other"])
                priority = st.selectbox("Priority", ["🔥 High", "⚡ Medium", "💫 Low"])
                
                col_add, col_cancel = st.columns(2)
                with col_add:
                    submitted = st.form_submit_button("Add Task")
                    if submitted and task_desc:
                        todos.append({
                            "id": max([todo.get('id', 0) for todo in todos], default=0) + 1,
                            "task": task_desc,
                            "category": category,
                            "priority": priority,
                            "due_date": selected_date.strftime("%Y-%m-%d"),
                            "completed": False,
                            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        })
                        save_todos(todos)
                        st.session_state.quick_add_for_date = None
                        st.success("Task added!")
                        st.rerun()
                
                with col_cancel:
                    if st.form_submit_button("Cancel"):
                        st.session_state.quick_add_for_date = None
                        st.rerun()
        
        # Close popup button
        if st.button("❌ Close", key="close_popup"):
            st.session_state.show_task_popup = False
            st.session_state.selected_date = None
            st.session_state.popup_tasks = []
            st.rerun()
    
    # Color coding legend
    st.markdown("---")
    st.markdown("### 🎨 Color Legend")
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("🔴 **Urgent (0-3 days)**")
    with col2:
        st.markdown("🟡 **Warning (4-6 days)**")
    with col3:
        st.markdown("🟢 **Safe (7+ days)**")
    with col4:
        st.markdown("🔵 **Today**")
    
    # Today's tasks summary
    st.markdown("---")
    today_tasks = [todo for todo in todos if todo['due_date'] == current_date.strftime("%Y-%m-%d")]
    
    if today_tasks:
        st.markdown("### 📋 Today's Tasks")
        
        for i, task in enumerate(today_tasks):
            col1, col2, col3 = st.columns([3, 1, 1])
            
            with col1:
                done = st.checkbox(
                    task["task"],
                    key=f"today_task_{i}",
                    value=task["completed"]
                )
                if done != task["completed"]:
                    for j, todo in enumerate(todos):
                        if todo['id'] == task['id']:
                            todos[j]["completed"] = done
                            break
                    save_todos(todos)
                    st.rerun()
            
            with col2:
                st.write(f"{task['category']} | {task['priority']}")
            
            with col3:
                if st.button("🗑️", key=f"delete_today_{i}"):
                    todos = [todo for todo in todos if todo['id'] != task['id']]
                    save_todos(todos)
                    st.rerun()
    else:
        st.info("🎉 No tasks due today! Great job staying organized!")
    
    # Quick stats with color coding
    st.markdown("### 📊 Quick Stats")
    col1, col2, col3, col4, col5 = st.columns(5)
    
    total_tasks = len([todo for todo in todos if not todo["completed"]])
    completed_tasks = len([todo for todo in todos if todo["completed"]])
    
    # Categorize by urgency
    urgent_tasks = len([todo for todo in todos if not todo["completed"] and get_urgency_class(get_days_until_deadline(todo['due_date'])) == "urgent"])
    warning_tasks = len([todo for todo in todos if not todo["completed"] and get_urgency_class(get_days_until_deadline(todo['due_date'])) == "warning"])
    safe_tasks = len([todo for todo in todos if not todo["completed"] and get_urgency_class(get_days_until_deadline(todo['due_date'])) == "safe"])
    
    with col1:
        st.metric("Active Tasks 📝", total_tasks)
    
    with col2:
        st.metric("Completed ✅", completed_tasks)
    
    with col3:
        st.metric("Urgent 🔴", urgent_tasks)
    
    with col4:
        st.metric("Warning 🟡", warning_tasks)
    
    with col5:
        st.metric("Safe 🟢", safe_tasks)




# === Grade Calculator //////////////////////////////////////////////////////////////////////////////////////===
def grade_calculator():
    st.header("📊 Grade Calculator")
    col1, col2 = st.columns(2)
    with col1:
        total_marks = st.number_input("Total Marks 📝", min_value=1, value=100)
    with col2:
        obtained_marks = st.number_input("Marks You Got ✍️", min_value=0, max_value=total_marks, value=0)
    if st.button("Calculate My Grade 🎯"):
        percent = obtained_marks / total_marks * 100
        if percent >= 90:
            grade = 'A'
        elif percent >= 80:
            grade = 'B'
        elif percent >= 70:
            grade = 'C'
        elif percent >= 60:
            grade = 'D'
        else:
            grade = 'F'
        emojis = {'A': '🌟', 'B': '✨', 'C': '👍', 'D': '😊', 'F': '💪'}
        st.markdown(f"### Results 📊\n- Percentage: **{percent:.2f}%**\n- Grade: **{grade}** {emojis[grade]}")





# === SGPA Calculator //////////////////////////////////////////////////////////////////////////////////////===
def sgpa_calculator():
    st.header("🎓 SGPA Calculator")
    num_subj = st.number_input("Number of Subjects 📚", min_value=1, max_value=10, value=5)
    total_credits = 0
    weighted_sum = 0.0
    st.markdown("### Enter Your Grades 📝")
    for i in range(int(num_subj)):
        colA, colB = st.columns(2)
        with colA:
            gp = st.number_input(f"Grade Points Subject {i+1} 📊", min_value=0.0, max_value=10.0, value=0.0)
        with colB:
            cr = st.number_input(f"Credits Subject {i+1} 💫", min_value=1, max_value=6, value=3)
        weighted_sum += gp * cr
        total_credits += cr
    if st.button("Calculate SGPA 🎯"):
        if total_credits:
            sgpa = weighted_sum / total_credits
            st.markdown(f"### Your SGPA ✨: **{sgpa:.2f}** 🎉")
        else:
            st.error("Total credits cannot be zero! 😅")





# === CGPA Calculator //////////////////////////////////////////////////////////////////////////////////////===
def cgpa_calculator():
    st.subheader("🏅 CGPA Calculator")
    sem_count = st.selectbox("Select Number of Semesters 📘", list(range(1, 9)), index=7)
    total_sgpa = 0.0
    st.markdown("### Enter SGPA for Each Semester 🎓")
    for i in range(sem_count):
        sgpa = st.number_input(f"SGPA for Semester {i+1}", min_value=0.0, max_value=10.0, value=0.0, key=f"sem_{i}")
        total_sgpa += sgpa
    if st.button("Calculate CGPA 🎯"):
        if sem_count:
            cgpa = total_sgpa / sem_count
            percentage = (cgpa - 0.75) * 10
            st.success(f"### Your CGPA Across {sem_count} Semesters is: {cgpa:.2f} 🌟")
            st.info(f"🎓 Estimated Percentage: {percentage:.2f}%")
        else:
            st.error("Number of semesters must be at least 1! 😅")




# === Physics Solver //////////////////////////////////////////////////////////////////////////////////////===
def physics_solver():
    st.header("⚡ Physics Formula Solver")
    """Basic physics problem solver"""
    st.header("⚡ Physics Solver")
    
    formula_type = st.selectbox("Choose formula type:", [
        "Velocity (v = u + at)",
        "Distance (s = ut + 0.5at²)",
        "Kinetic Energy (KE = 0.5mv²)",
        "Potential Energy (PE = mgh)"
    ])
    
    if formula_type == "Velocity (v = u + at)":
        u = st.number_input("Initial velocity (u):", value=0.0)
        a = st.number_input("Acceleration (a):", value=0.0)
        t = st.number_input("Time (t):", value=0.0)
        
        if st.button("Calculate"):
            v = u + a * t
            st.success(f"Final velocity (v) = {v} m/s")
    
    elif formula_type == "Distance (s = ut + 0.5at²)":
        u = st.number_input("Initial velocity (u):", value=0.0)
        t = st.number_input("Time (t):", value=0.0)
        a = st.number_input("Acceleration (a):", value=0.0)
        
        if st.button("Calculate"):
            s = u * t + 0.5 * a * t * t
            st.success(f"Distance (s) = {s} m")
    
    elif formula_type == "Kinetic Energy (KE = 0.5mv²)":
        m = st.number_input("Mass (m):", value=0.0)
        v = st.number_input("Velocity (v):", value=0.0)
        
        if st.button("Calculate"):
            ke = 0.5 * m * v * v
            st.success(f"Kinetic Energy (KE) = {ke} J")
    
    elif formula_type == "Potential Energy (PE = mgh)":
        m = st.number_input("Mass (m):", value=0.0)
        g = st.number_input("Gravity (g):", value=9.8)
        h = st.number_input("Height (h):", value=0.0)
        
        if st.button("Calculate"):
            pe = m * g * h
            st.success(f"Potential Energy (PE) = {pe} J")





# === Unit Converter //////////////////////////////////////////////////////////////////////////////////////===
def unit_converter():
    st.header("📏 Unit Converter")
    """Unit converter for various measurements"""
    st.header("📏 Unit Converter")
    
    conversion_type = st.selectbox("Choose conversion type:", [
        "Length", "Weight", "Temperature", "Area", "Volume"
    ])
    
    if conversion_type == "Length":
        value = st.number_input("Enter value:")
        from_unit = st.selectbox("From:", ["meters", "kilometers", "centimeters", "feet", "inches"])
        to_unit = st.selectbox("To:", ["meters", "kilometers", "centimeters", "feet", "inches"])
        
        if st.button("Convert"):
            # Convert to meters first
            if from_unit == "kilometers":
                meters = value * 1000
            elif from_unit == "centimeters":
                meters = value / 100
            elif from_unit == "feet":
                meters = value * 0.3048
            elif from_unit == "inches":
                meters = value * 0.0254
            else:
                meters = value
            
            # Convert from meters to target unit
            if to_unit == "kilometers":
                result = meters / 1000
            elif to_unit == "centimeters":
                result = meters * 100
            elif to_unit == "feet":
                result = meters / 0.3048
            elif to_unit == "inches":
                result = meters / 0.0254
            else:
                result = meters
            
            st.success(f"{value} {from_unit} = {result:.4f} {to_unit}")




# === Inline URL Shortener //////////////////////////////////////////////////////////////////////////////////////===
def url_shortener_page():
    st.header("🔗 URL Shortener")
    long_url = st.text_input("Paste URL here to shorten:")
    if st.button("Shorten URL"):
        if not long_url:
            st.error("Please enter a URL to shorten.")
        else:
            try:
                shortener = pyshorteners.Shortener(timeout=5)
                try:
                    result = shortener.tinyurl.short(long_url)
                except:
                    try:
                        result = shortener.isgd.short(long_url)
                    except Exception:
                        result = shortener.dagd.short(long_url)
                st.success("URL shortened successfully!")
                st.write(f"🔗 {result}")
                # Automatically copy to clipboard
                st.write("(Copied to clipboard)")
            except Exception as e:
                st.error(f"Error: {e}")




# === Flashcard Generator Launcher ///////////////////////////////////////////////////////////////////////////////////////===
def generate_random_color():
    """Generate a random pastel color that works well with text."""
    h = random.random()  # Hue (0-1)
    s = 0.7 + random.random() * 0.2  # Saturation (0.7-0.9)
    l = 0.65 + random.random() * 0.15  # Lightness (0.65-0.8)
    
    # Convert HSL to RGB
    if s == 0:
        r = g = b = l
    else:
        def hue_to_rgb(p, q, t):
            if t < 0: t += 1
            if t > 1: t -= 1
            if t < 1/6: return p + (q - p) * 6 * t
            if t < 1/2: return q
            if t < 2/3: return p + (q - p) * (2/3 - t) * 6
            return p
        
        q = l * (1 + s) if l < 0.5 else l + s - l * s
        p = 2 * l - q
        r = hue_to_rgb(p, q, h + 1/3)
        g = hue_to_rgb(p, q, h)
        b = hue_to_rgb(p, q, h - 1/3)
    
    # Convert to hex
    r, g, b = [int(x * 255) for x in (r, g, b)]
    return f"#{r:02x}{g:02x}{b:02x}"

def record_voice_note():
    """Record voice note using browser's audio recording capabilities."""
    st.markdown("### 🎤 Voice Note Recording")
    
    # JavaScript for audio recording
    audio_recorder_js = """
    <div id="audio-recorder">
        <button id="startBtn" onclick="startRecording()">🎤 Start Recording</button>
        <button id="stopBtn" onclick="stopRecording()" disabled>⏹️ Stop Recording</button>
        <button id="playBtn" onclick="playRecording()" disabled>▶️ Play Recording</button>
        <div id="status">Ready to record</div>
        <audio id="audioPlayback" controls style="display:none;"></audio>
    </div>

    <script>
    let mediaRecorder;
    let audioChunks = [];
    let audioBlob;

    async function startRecording() {
        try {
            const stream = await navigator.mediaDevices.getUserMedia({ audio: true });
            mediaRecorder = new MediaRecorder(stream);
            
            mediaRecorder.ondataavailable = event => {
                audioChunks.push(event.data);
            };
            
            mediaRecorder.onstop = () => {
                audioBlob = new Blob(audioChunks, { type: 'audio/wav' });
                const audioUrl = URL.createObjectURL(audioBlob);
                document.getElementById('audioPlayback').src = audioUrl;
                document.getElementById('audioPlayback').style.display = 'block';
                document.getElementById('playBtn').disabled = false;
                
                // Convert to base64 and send to Streamlit
                const reader = new FileReader();
                reader.onloadend = function() {
                    const base64data = reader.result;
                    window.parent.postMessage({
                        type: 'audio_recorded',
                        data: base64data
                    }, '*');
                };
                reader.readAsDataURL(audioBlob);
            };
            
            audioChunks = [];
            mediaRecorder.start();
            
            document.getElementById('startBtn').disabled = true;
            document.getElementById('stopBtn').disabled = false;
            document.getElementById('status').innerHTML = '🔴 Recording...';
            
        } catch (err) {
            console.error('Error accessing microphone:', err);
            document.getElementById('status').innerHTML = '❌ Error: Could not access microphone';
        }
    }

    function stopRecording() {
        if (mediaRecorder && mediaRecorder.state === 'recording') {
            mediaRecorder.stop();
            mediaRecorder.stream.getTracks().forEach(track => track.stop());
            
            document.getElementById('startBtn').disabled = false;
            document.getElementById('stopBtn').disabled = true;
            document.getElementById('status').innerHTML = '✅ Recording complete!';
        }
    }

    function playRecording() {
        const audio = document.getElementById('audioPlayback');
        audio.play();
    }
    </script>

    <style>
    #audio-recorder {
        padding: 20px;
        border: 2px dashed #ccc;
        border-radius: 10px;
        text-align: center;
        margin: 10px 0;
    }
    
    #audio-recorder button {
        margin: 5px;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        font-size: 14px;
    }
    
    #startBtn { background-color: #4CAF50; color: white; }
    #stopBtn { background-color: #f44336; color: white; }
    #playBtn { background-color: #2196F3; color: white; }
    
    #status {
        margin-top: 10px;
        font-weight: bold;
    }
    </style>
    """
    
    st.components.v1.html(audio_recorder_js, height=200)
    
    # Alternative: File upload for audio
    st.markdown("**Or upload an audio file:**")
    uploaded_audio = st.file_uploader(
        "Upload audio file", 
        type=['mp3', 'wav', 'm4a', 'ogg', 'flac'],
        key="voice_note_upload"
    )
    
    return uploaded_audio

def extract_text_from_audio(audio_file, language='en-US'):
    """Extract text from audio file using speech recognition."""
    text = ""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
            temp_filename = temp_file.name
            
        # Convert audio to WAV format if needed
        if hasattr(audio_file, 'name') and audio_file.name.lower().endswith(('.mp3', '.m4a', '.ogg', '.flac')):
            audio = AudioSegment.from_file(io.BytesIO(audio_file.read()))
            audio.export(temp_filename, format="wav")
        else:
            # Assume it's already in WAV format
            with open(temp_filename, 'wb') as f:
                f.write(audio_file.read() if hasattr(audio_file, 'read') else audio_file)
        
        # Initialize recognizer
        recognizer = sr.Recognizer()
        
        # Load audio file
        with sr.AudioFile(temp_filename) as source:
            # Adjust for ambient noise
            recognizer.adjust_for_ambient_noise(source)
            audio_data = recognizer.record(source)
            
        # Recognize speech
        text = recognizer.recognize_google(audio_data, language=language)
        
        # Clean up temporary file
        os.unlink(temp_filename)
        
    except sr.UnknownValueError:
        st.error("Could not understand audio")
    except sr.RequestError as e:
        st.error(f"Could not request results from speech recognition service: {e}")
    except Exception as e:
        st.error(f"Error processing audio: {e}")
        # Clean up temporary file if it exists
        if 'temp_filename' in locals() and os.path.exists(temp_filename):
            os.unlink(temp_filename)
    
    return text

def create_flashcards_from_voice_notes():
    """Create flashcards specifically from voice notes with intelligent parsing."""
    st.subheader("🎙️ Voice Note to Flashcards")
    
    # Voice recording section
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("**Record or upload your voice notes:**")
        uploaded_audio = record_voice_note()
    
    with col2:
        # Language selection for voice recognition
        st.markdown("**Language Settings:**")
        language_options = {
            'English (US)': 'en-US',
            'English (UK)': 'en-GB',
            'Spanish': 'es-ES',
            'French': 'fr-FR',
            'German': 'de-DE',
            'Italian': 'it-IT',
            'Portuguese': 'pt-PT',
            'Russian': 'ru-RU',
            'Chinese': 'zh-CN',
            'Japanese': 'ja-JP'
        }
        selected_language = st.selectbox("Select language:", list(language_options.keys()))
        language = language_options[selected_language]
    
    # Voice note instructions
    with st.expander("💡 Tips for Better Voice Note Flashcards"):
        st.markdown("""
        **For best results, structure your voice notes like this:**
        
        📝 **Question-Answer Format:**
        - "Question: What is photosynthesis? Answer: The process by which plants convert sunlight into energy."
        
        📝 **Definition Format:**
        - "Photosynthesis is the process by which plants convert sunlight into energy."
        
        📝 **List Format:**
        - "The three types of rocks are igneous, sedimentary, and metamorphic."
        
        📝 **Study Points:**
        - Speak clearly and at a moderate pace
        - Pause between different topics
        - Use keywords like "Question", "Answer", "Definition", "Remember"
        - Repeat important terms for clarity
        """)
    
    # Process voice notes
    if uploaded_audio:
        with st.spinner("🎧 Processing voice note and creating flashcards..."):
            # Extract text from audio
            transcribed_text = extract_text_from_audio(uploaded_audio, language)
            
            if transcribed_text:
                st.success("✅ Voice note transcribed successfully!")
                
                # Show transcribed text
                with st.expander("📝 View Transcribed Text"):
                    st.text_area("Transcription:", transcribed_text, height=150)
                
                # Create smart flashcards from voice notes
                voice_flashcards = create_smart_voice_flashcards(transcribed_text)
                
                if voice_flashcards:
                    # Store in session state
                    if 'flashcards' not in st.session_state:
                        st.session_state['flashcards'] = []
                        st.session_state['flashcard_colors'] = []
                    
                    # Add new flashcards
                    st.session_state['flashcards'].extend(voice_flashcards)
                    st.session_state['flashcard_colors'].extend([generate_random_color() for _ in voice_flashcards])
                    
                    st.success(f"🎉 Created {len(voice_flashcards)} flashcards from your voice note!")
                    
                    # Show preview of created flashcards
                    with st.expander("👀 Preview Generated Flashcards"):
                        for i, card in enumerate(voice_flashcards):
                            st.markdown(f"**Card {i+1}:**")
                            st.info(f"Q: {card['question']}")
                            st.success(f"A: {card['answer']}")
                            st.markdown("---")
                
                else:
                    st.warning("Could not create flashcards from the voice note. Try speaking more clearly or using the suggested formats.")
            
            else:
                st.error("Could not transcribe the voice note. Please try again with clearer audio.")

# def create_flashcards_from_voice_notes():
#     """Create flashcards specifically from voice notes with intelligent parsing."""
#     st.subheader("🎙️ Voice Note to Flashcards")
    
#     # Voice recording section
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         st.markdown("**Record or upload your voice notes:**")
#         uploaded_audio = record_voice_note()
    
#     with col2:
#         # Language selection for voice recognition
#         st.markdown("**Language Settings:**")
#         language_options = {
#             'English (US)': 'en-US',
#             'English (UK)': 'en-GB',
#             'Spanish': 'es-ES',
#             'French': 'fr-FR',
#             'German': 'de-DE',
#             'Italian': 'it-IT',
#             'Portuguese': 'pt-PT',
#             'Russian': 'ru-RU',
#             'Chinese': 'zh-CN',
#             'Japanese': 'ja-JP'
#         }
#         selected_language = st.selectbox("Select language:", list(language_options.keys()))
#         language = language_options[selected_language]
    
#     # Voice note instructions
#     with st.expander("💡 Tips for Better Voice Note Flashcards"):
#         st.markdown("""
#         **For best results, structure your voice notes like this:**
        
#         📝 **Question-Answer Format:**
#         - "Question: What is photosynthesis? Answer: The process by which plants convert sunlight into energy."
        
#         📝 **Definition Format:**
#         - "Photosynthesis is the process by which plants convert sunlight into energy."
        
#         📝 **List Format:**
#         - "The three types of rocks are igneous, sedimentary, and metamorphic."
        
#         📝 **Study Points:**
#         - Speak clearly and at a moderate pace
#         - Pause between different topics
#         - Use keywords like "Question", "Answer", "Definition", "Remember"
#         - Repeat important terms for clarity
#         """)
    
#     # Process voice notes
#     if uploaded_audio:
#         with st.spinner("🎧 Processing voice note and creating flashcards..."):
#             # Extract text from audio
#             transcribed_text = extract_text_from_audio(uploaded_audio, language)
            
#             if transcribed_text:
#                 st.success("✅ Voice note transcribed successfully!")
                
#                 # Show transcribed text
#                 with st.expander("📝 View Transcribed Text"):
#                     st.text_area("Transcription:", transcribed_text, height=150)
                
#                 # Create smart flashcards from voice notes
#                 voice_flashcards = create_smart_voice_flashcards(transcribed_text)
                
#                 if voice_flashcards:
#                     # Store in session state
#                     if 'flashcards' not in st.session_state:
#                         st.session_state['flashcards'] = []
#                         st.session_state['flashcard_colors'] = []
                    
#                     # Add new flashcards
#                     st.session_state['flashcards'].extend(voice_flashcards)
#                     st.session_state['flashcard_colors'].extend([generate_random_color() for _ in voice_flashcards])
                    
#                     st.success(f"🎉 Created {len(voice_flashcards)} flashcards from your voice note!")
                    
#                     # Show preview of created flashcards
#                     with st.expander("👀 Preview Generated Flashcards"):
#                         for i, card in enumerate(voice_flashcards):
#                             st.markdown(f"**Card {i+1}:**")
#                             st.info(f"Q: {card['question']}")
#                             st.success(f"A: {card['answer']}")
#                             st.markdown("---")
                
#                 else:
#                     st.warning("Could not create flashcards from the voice note. Try speaking more clearly or using the suggested formats.")
            
#             else:
#                 st.error("Could not transcribe the voice note. Please try again with clearer audio.")

def create_smart_voice_flashcards(text):
    """Create intelligent flashcards from voice note transcriptions."""
    flashcards = []
    
    # Clean and prepare text
    text = text.strip()
    
    # Method 1: Direct Question-Answer patterns
    qa_patterns = [
        r'question[:\s]+([^?]+\?)\s*answer[:\s]+([^.!?]+[.!?])',
        r'q[:\s]+([^?]+\?)\s*a[:\s]+([^.!?]+[.!?])',
        r'([^.!?]+\?)\s*([^.!?]+[.!?])',  # Simple Q? A. pattern
    ]
    
    for pattern in qa_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        for question, answer in matches:
            question = question.strip()
            answer = answer.strip()
            if question and answer and len(answer) > 5:
                flashcards.append({
                    "question": question if question.endswith('?') else question + '?',
                    "answer": answer
                })
    
    # Method 2: Definition patterns
    definition_patterns = [
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+is\s+([^.!?]+[.!?])',
        r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+means\s+([^.!?]+[.!?])',
        r'define\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)[:\s]+([^.!?]+[.!?])',
        r'remember\s+that\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+([^.!?]+[.!?])',
    ]
    
    for pattern in definition_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        for term, definition in matches:
            term = term.strip()
            definition = definition.strip()
            if term and definition and len(definition) > 10:
                question = f"What is {term}?"
                flashcards.append({
                    "question": question,
                    "answer": definition
                })
    
    # Method 3: List patterns
    list_patterns = [
        r'(?:the\s+)?(?:three|four|five|six|seven|eight|nine|ten|\d+)\s+(?:types?|kinds?|categories|examples?)\s+of\s+([^.!?]+)\s+are\s+([^.!?]+[.!?])',
        r'([^.!?]+)\s+(?:include|are|consists?\s+of)[:\s]+([^.!?]+[.!?])',
    ]
    
    for pattern in list_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        for topic, items in matches:
            topic = topic.strip()
            items = items.strip()
            if topic and items and len(items) > 10:
                question = f"What are the types/examples of {topic}?"
                flashcards.append({
                    "question": question,
                    "answer": items
                })
    
    # Method 4: Important facts pattern
    fact_patterns = [
        r'(?:remember|note|important)[:\s]+([^.!?]+[.!?])',
        r'(?:fact|key\s+point)[:\s]+([^.!?]+[.!?])',
    ]
    
    for pattern in fact_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
        for fact in matches:
            fact = fact.strip()
            if fact and len(fact) > 15:
                # Create a question from the fact
                first_few_words = ' '.join(fact.split()[:4])
                question = f"What should you remember about {first_few_words}...?"
                flashcards.append({
                    "question": question,
                    "answer": fact
                })
    
    # Method 5: Sentence splitting for general content
    if len(flashcards) < 3:  # If we don't have enough flashcards, try sentence approach
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]
        
        for i in range(0, len(sentences) - 1, 2):  # Take pairs of sentences
            if i + 1 < len(sentences):
                statement = sentences[i].strip()
                context = sentences[i + 1].strip()
                
                if statement and context:
                    # Create question from statement
                    question = f"Complete this statement: {statement[:30]}...?"
                    answer = f"{statement}. {context}"
                    flashcards.append({
                        "question": question,
                        "answer": answer
                    })
    
    # Remove duplicates and clean up
    unique_flashcards = []
    seen_questions = set()
    
    for card in flashcards:
        # Clean text
        card['question'] = re.sub(r'\s+', ' ', card['question']).strip()
        card['answer'] = re.sub(r'\s+', ' ', card['answer']).strip()
        
        # Check for duplicates
        question_key = card['question'].lower()
        if question_key not in seen_questions and len(card['answer']) > 5:
            seen_questions.add(question_key)
            unique_flashcards.append(card)
    
    return unique_flashcards[:20]  # Limit to 20 flashcards


def extract_images_from_pdf(file_stream):
    """Extract images from PDF using PyMuPDF."""
    images = []
    try:
        # Read PDF with PyMuPDF
        file_stream.seek(0)
        pdf_document = fitz.open(stream=file_stream.read(), filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                # Get image data
                xref = img[0]
                pix = fitz.Pixmap(pdf_document, xref)
                
                # Convert to PIL Image
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("pil")
                    img_io = io.BytesIO(img_data)
                    pil_image = Image.open(img_io)
                    
                    # Convert to RGB if necessary
                    if pil_image.mode in ('RGBA', 'LA', 'P'):
                        pil_image = pil_image.convert('RGB')
                    
                    # Resize if too large
                    max_size = (800, 600)
                    pil_image.thumbnail(max_size, Image.Resampling.LANCZOS)
                    
                    # Convert to base64
                    img_buffer = BytesIO()
                    pil_image.save(img_buffer, format='JPEG', quality=85)
                    img_str = base64.b64encode(img_buffer.getvalue()).decode()
                    images.append(img_str)
                
                pix = None  # Free memory
        
        pdf_document.close()
        
    except Exception as e:
        st.warning(f"Could not extract images from PDF: {e}")
        # Fallback to basic PDF processing without images
    
    return images

def extract_images_from_docx(file_stream):
    """Extract images from DOCX file."""
    images = []
    try:
        file_stream.seek(0)
        # Read DOCX as ZIP file to access images
        with zipfile.ZipFile(file_stream, 'r') as docx_zip:
            # List all files in the DOCX
            file_list = docx_zip.namelist()
            
            # Find image files
            image_files = [f for f in file_list if f.startswith('word/media/') and 
                          f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
            
            for image_file in image_files:
                try:
                    # Extract image data
                    img_data = docx_zip.read(image_file)
                    img_io = io.BytesIO(img_data)
                    pil_image = Image.open(img_io)
                    
                    # Convert to RGB if necessary
                    if pil_image.mode in ('RGBA', 'LA', 'P'):
                        pil_image = pil_image.convert('RGB')
                    
                    # Resize if too large
                    max_size = (800, 600)
                    pil_image.thumbnail(max_size, Image.Resampling.LANCZOS)
                    
                    # Convert to base64
                    img_buffer = BytesIO()
                    pil_image.save(img_buffer, format='JPEG', quality=85)
                    img_str = base64.b64encode(img_buffer.getvalue()).decode()
                    images.append(img_str)
                    
                except Exception as e:
                    st.warning(f"Could not process image {image_file}: {e}")
                    continue
    
    except Exception as e:
        st.warning(f"Could not extract images from DOCX: {e}")
    
    return images

def extract_text_from_pdf(file_stream):
    """Extract text from a PDF file."""
    text = ""
    try:
        file_stream.seek(0)
        pdf_reader = PdfReader(file_stream)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file_stream):
    """Extract text from a Word document."""
    text = ""
    try:
        file_stream.seek(0)
        doc = Document(file_stream)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
    return text

def extract_text_from_audio(audio_file, language='en-US'):
    """Extract text from audio file using speech recognition."""
    text = ""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
            temp_filename = temp_file.name
            
        # Convert audio to WAV format if needed
        if audio_file.name.lower().endswith(('.mp3', '.m4a', '.ogg', '.flac')):
            audio = AudioSegment.from_file(io.BytesIO(audio_file.read()))
            audio.export(temp_filename, format="wav")
        else:
            # Assume it's already in WAV format
            with open(temp_filename, 'wb') as f:
                f.write(audio_file.read())
        
        # Initialize recognizer
        recognizer = sr.Recognizer()
        
        # Load audio file
        with sr.AudioFile(temp_filename) as source:
            # Adjust for ambient noise
            recognizer.adjust_for_ambient_noise(source)
            audio_data = recognizer.record(source)
            
        # Recognize speech
        text = recognizer.recognize_google(audio_data, language=language)
        
        # Clean up temporary file
        os.unlink(temp_filename)
        
    except sr.UnknownValueError:
        st.error("Could not understand audio")
    except sr.RequestError as e:
        st.error(f"Could not request results from speech recognition service: {e}")
    except Exception as e:
        st.error(f"Error processing audio: {e}")
        # Clean up temporary file if it exists
        if 'temp_filename' in locals() and os.path.exists(temp_filename):
            os.unlink(temp_filename)
    
    return text

def process_image_to_base64(image_file):
    """Convert uploaded image to base64 string for storage."""
    try:
        image = Image.open(image_file)
        # Resize image if too large
        max_size = (800, 600)
        image.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Convert to RGB if necessary
        if image.mode in ('RGBA', 'LA', 'P'):
            image = image.convert('RGB')
        
        # Save to bytes
        img_buffer = BytesIO()
        image.save(img_buffer, format='JPEG', quality=85)
        img_str = base64.b64encode(img_buffer.getvalue()).decode()
        
        return img_str
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None

def extract_questions_and_answers(text):
    """Extract potential question-answer pairs from text."""
    flashcards = []
    
    # Method 1: Look for question marks followed by text
    question_pattern = re.compile(r'([^.!?]+\?)\s*([^?]+?)(?=\s*[A-Z]|$)', re.DOTALL)
    matches = question_pattern.findall(text)
    
    for question, answer in matches:
        question = question.strip()
        answer = answer.strip()
        if question and answer and len(answer) > 5:
            flashcards.append({"question": question, "answer": answer})
    
    # Method 2: Look for numbered or bulleted items
    numbered_pattern = re.compile(r'(\d+[.)]\s*|-\s+|\*\s+)([^.!?]+[.!?])\s*([^.!?]+[.!?])', re.DOTALL)
    matches = numbered_pattern.findall(text)
    
    for _, potential_question, potential_answer in matches:
        question = potential_question.strip()
        answer = potential_answer.strip()
        
        if question and answer and "?" not in question:
            if not question.endswith("?"):
                question = question[:-1] + "?" if question.endswith((".", "!")) else question + "?"
            flashcards.append({"question": question, "answer": answer})
    
    # Method 3: Look for definition patterns (term: definition)
    definition_pattern = re.compile(r'([A-Z][^:]+):\s*([^.!?]+[.!?])', re.MULTILINE)
    matches = definition_pattern.findall(text)
    
    for term, definition in matches:
        term = term.strip()
        definition = definition.strip()
        if term and definition and len(definition) > 10:
            question = f"What is {term}?"
            flashcards.append({"question": question, "answer": definition})
    
    # If we couldn't extract enough flashcards, create some from paragraphs
    if len(flashcards) < 5:
        paragraphs = re.split(r'\n\s*\n', text)
        
        for i in range(len(paragraphs) - 1):
            paragraph = paragraphs[i].strip()
            next_paragraph = paragraphs[i + 1].strip()
            
            if len(paragraph) > 20 and len(next_paragraph) > 20:
                sentences = re.split(r'[.!?]+', paragraph)
                sentences = [s.strip() for s in sentences if s.strip()]
                
                if sentences:
                    main_sentence = sentences[0]
                    preview = main_sentence[:30] + "..." if len(main_sentence) > 30 else main_sentence
                    question = f'What is the main point about "{preview}"?'
                    flashcards.append({"question": question, "answer": paragraph})
    
    # Deduplicate flashcards
    unique_flashcards = []
    seen_questions = set()
    
    for card in flashcards:
        if card["question"] not in seen_questions:
            seen_questions.add(card["question"])
            unique_flashcards.append(card)
    
    return unique_flashcards[:50]  # Limit to 50 flashcards

def clean_text(text):
    """Clean text by removing HTML tags and extra whitespace."""
    # Remove HTML tags
    clean = re.sub('<.*?>', '', text)
    # Remove extra whitespace
    clean = re.sub(r'\s+', ' ', clean)
    return clean.strip()

def save_flashcards_to_file(flashcards, colors, images=None, filename=None):
    """Save flashcards to a JSON file."""
    if filename is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"flashcards_{timestamp}.json"
    
    data = {
        'flashcards': flashcards,
        'colors': colors,
        'images': images or [],
        'created_at': datetime.now().isoformat(),
        'total_cards': len(flashcards)
    }
    
    try:
        # Create flashcards directory if it doesn't exist
        os.makedirs('saved_flashcards', exist_ok=True)
        filepath = os.path.join('saved_flashcards', filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return filepath
    except Exception as e:
        st.error(f"Error saving flashcards: {e}")
        return None

def load_flashcards_from_file(filepath):
    """Load flashcards from a JSON file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except Exception as e:
        st.error(f"Error loading flashcards: {e}")
        return None

def get_saved_flashcard_files():
    """Get list of saved flashcard files."""
    try:
        if not os.path.exists('saved_flashcards'):
            return []
        
        files = []
        for filename in os.listdir('saved_flashcards'):
            if filename.endswith('.json'):
                filepath = os.path.join('saved_flashcards', filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    files.append({
                        'filename': filename,
                        'filepath': filepath,
                        'created_at': data.get('created_at', 'Unknown'),
                        'total_cards': data.get('total_cards', 0)
                    })
                except:
                    continue
        
        return sorted(files, key=lambda x: x['created_at'], reverse=True)
    except Exception as e:
        st.error(f"Error listing saved flashcards: {e}")
        return []

def display_flashcard(card, color, image_data=None, card_index=0):
    """Display a single flashcard with flip functionality."""
    card_key = f"card_{card_index}"
    
    # Initialize flip state if not exists
    if f"flipped_{card_key}" not in st.session_state:
        st.session_state[f"flipped_{card_key}"] = False
    
    # Clean the text content
    clean_question = clean_text(card['question'])
    clean_answer = clean_text(card['answer'])
    
    # Create card container
    with st.container():
        col1, col2, col3 = st.columns([1, 8, 1])
        
        with col2:
            # Create a styled container using Streamlit's native components
            if not st.session_state[f"flipped_{card_key}"]:
                # Question side
                with st.container():
                    st.markdown(f"""
                    <div style="
                        background-color: {color};
                        border-radius: 15px;
                        padding: 30px;
                        margin: 15px 0;
                        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
                        border: 1px solid rgba(0,0,0,0.1);
                        min-height: 200px;
                        text-align: center;
                        display: flex;
                        flex-direction: column;
                        justify-content: center;
                    ">
                        <h3 style="color: #333; margin-bottom: 20px; font-size: 24px;">❓ Question</h3>
                        <p style="color: #555; font-size: 18px; line-height: 1.6; margin: 0;">{clean_question}</p>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                # Answer side
                with st.container():
                    st.markdown(f"""
                    <div style="
                        background-color: {color};
                        border-radius: 15px;
                        padding: 30px;
                        margin: 15px 0;
                        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
                        border: 1px solid rgba(0,0,0,0.1);
                        min-height: 200px;
                        text-align: center;
                        display: flex;
                        flex-direction: column;
                        justify-content: center;
                    ">
                        <h3 style="color: #333; margin-bottom: 20px; font-size: 24px;">💡 Answer</h3>
                        <p style="color: #555; font-size: 16px; line-height: 1.6; margin: 0;">{clean_answer}</p>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Display image if available
            if image_data:
                try:
                    img_bytes = base64.b64decode(image_data)
                    image = Image.open(io.BytesIO(img_bytes))
                    st.image(image, width=300, caption="Reference Image")
                except Exception as e:
                    st.error(f"Error displaying image: {e}")
            
            # Flip button
            st.markdown("<br>", unsafe_allow_html=True)
            col_a, col_b, col_c = st.columns([2, 1, 2])
            with col_b:
                if st.button("🔄 Flip Card", key=f"flip_{card_key}", use_container_width=True):
                    st.session_state[f"flipped_{card_key}"] = not st.session_state[f"flipped_{card_key}"]
                    st.rerun()

def add_flashcard_storage_section():
    """Add the flashcard storage and loading section."""
    st.subheader("💾 Save & Load Flashcards")
    
    # Save current flashcards
    if 'flashcards' in st.session_state and st.session_state['flashcards']:
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Save Current Flashcards**")
            custom_filename = st.text_input("Custom filename (optional)", placeholder="my_study_set")
            
            if st.button("💾 Save Flashcards", use_container_width=True):
                filename = f"{custom_filename}.json" if custom_filename else None
                filepath = save_flashcards_to_file(
                    st.session_state['flashcards'],
                    st.session_state['flashcard_colors'],
                    st.session_state.get('flashcard_images', []),
                    filename
                )
                if filepath:
                    st.success(f"✅ Flashcards saved to: {filepath}")
        
        with col2:
            st.markdown("**Quick Export**")
            export_data = {
                'flashcards': st.session_state['flashcards'],
                'colors': st.session_state['flashcard_colors'],
                'images': st.session_state.get('flashcard_images', []),
                'created_at': datetime.now().isoformat()
            }
            json_str = json.dumps(export_data, indent=2)
            st.download_button(
                label="📥 Download JSON",
                data=json_str,
                file_name=f"flashcards_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
    
    # Load saved flashcards
    st.markdown("**Load Saved Flashcards**")
    saved_files = get_saved_flashcard_files()
    
    if saved_files:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            selected_file = st.selectbox(
                "Select a saved flashcard set:",
                options=saved_files,
                format_func=lambda x: f"{x['filename']} ({x['total_cards']} cards) - {x['created_at'][:10]}",
                key="file_selector"
            )
        
        with col2:
            st.markdown("<br>", unsafe_allow_html=True)  # Add spacing
            if st.button("📂 Load", use_container_width=True):
                if selected_file:
                    data = load_flashcards_from_file(selected_file['filepath'])
                    if data:
                        st.session_state['flashcards'] = data['flashcards']
                        st.session_state['flashcard_colors'] = data['colors']
                        st.session_state['flashcard_images'] = data.get('images', [])
                        st.session_state['current_card'] = 0
                        st.success(f"✅ Loaded {len(data['flashcards'])} flashcards!")
                        st.rerun()
        
        # Display saved files info
        with st.expander("📁 View All Saved Sets"):
            for file_info in saved_files:
                col1, col2, col3 = st.columns([2, 1, 1])
                with col1:
                    st.text(file_info['filename'])
                with col2:
                    st.text(f"{file_info['total_cards']} cards")
                with col3:
                    st.text(file_info['created_at'][:10])
    else:
        st.info("No saved flashcard sets found. Create and save some flashcards first!")
    
    # Upload JSON file
    st.markdown("**Import Flashcards from File**")
    uploaded_json = st.file_uploader("Upload a flashcard JSON file", type=['json'])
    
    if uploaded_json:
        try:
            data = json.load(uploaded_json)
            if 'flashcards' in data and data['flashcards']:
                st.session_state['flashcards'] = data['flashcards']
                st.session_state['flashcard_colors'] = data.get('colors', [generate_random_color() for _ in data['flashcards']])
                st.session_state['flashcard_images'] = data.get('images', [])
                st.session_state['current_card'] = 0
                st.success(f"✅ Successfully imported {len(data['flashcards'])} flashcards!")
                st.rerun()
            else:
                st.error("Invalid flashcard file format.")
        except Exception as e:
            st.error(f"Error importing flashcards: {e}")

# Add this function to your main flashcard_generator_page() function
# Place it after the flashcard display section and before the usage instructions

# Alternative: Pure Streamlit components version (no custom HTML at all)
def display_flashcard_pure_streamlit(card, color, image_data=None, card_index=0):
    """Display flashcard using only native Streamlit components."""
    card_key = f"card_{card_index}"
    
    # Initialize flip state if not exists
    if f"flipped_{card_key}" not in st.session_state:
        st.session_state[f"flipped_{card_key}"] = False
    
    # Clean the text content
    clean_question = clean_text(card['question'])
    clean_answer = clean_text(card['answer'])
    
    # Create card using Streamlit container and styling
    with st.container():
        # Add some spacing
        st.markdown("---")
        
        if not st.session_state[f"flipped_{card_key}"]:
            # Question side
            st.markdown("### ❓ Question")
            st.info(clean_question)
        else:
            # Answer side  
            st.markdown("### 💡 Answer")
            st.success(clean_answer)
        
        # Display image if available
        if image_data:
            try:
                img_bytes = base64.b64decode(image_data)
                image = Image.open(io.BytesIO(img_bytes))
                st.image(image, width=300, caption="Reference Image")
            except Exception as e:
                st.error(f"Error displaying image: {e}")
        
        # Flip button
        col1, col2, col3 = st.columns([2, 1, 2])
        with col2:
            if st.button("🔄 Flip Card", key=f"flip_{card_key}", use_container_width=True):
                st.session_state[f"flipped_{card_key}"] = not st.session_state[f"flipped_{card_key}"]
                st.rerun()
        
        st.markdown("---")

def flashcard_generator_page():
    """Main flashcard generator page."""
    st.header("🃏 Enhanced Flashcard Generator")
    st.write("Generate flashcards from text files, documents, audio files, voice recordings, and optionally include images!")
    
    # Create tabs for different input methods
    tab1, tab2 = st.tabs(["📁 Upload Files", "🎤 Record Voice"])
    
    with tab1:
        # File upload section
        st.subheader("📁 Upload Content")
        
        # Main content upload
        uploaded_file = st.file_uploader(
            "Choose a file (PDF, DOCX, TXT, MP3, WAV, M4A, OGG, FLAC)",
            type=['pdf', 'docx', 'txt', 'mp3', 'wav', 'm4a', 'ogg', 'flac'],
            help="Upload text documents or audio files to generate flashcards",
            key="file_uploader"
        )
        
        # Optional image upload
        st.subheader("🖼️ Optional: Add Reference Images")
        uploaded_images = st.file_uploader(
            "Upload reference images (optional)",
            type=['png', 'jpg', 'jpeg', 'gif', 'bmp'],
            accept_multiple_files=True,
            help="Add images that will be displayed with your flashcards",
            key="image_uploader"
        )
        
        # Audio language selection (only show if audio file is uploaded)
        language = 'en-US'
        if uploaded_file and uploaded_file.name.lower().endswith(('.mp3', '.wav', '.m4a', '.ogg', '.flac')):
            st.subheader("🗣️ Audio Settings")
            language_options = {
                'English (US)': 'en-US',
                'English (UK)': 'en-GB',
                'Spanish': 'es-ES',
                'French': 'fr-FR',
                'German': 'de-DE',
                'Italian': 'it-IT',
                'Portuguese': 'pt-PT',
                'Russian': 'ru-RU',
                'Chinese': 'zh-CN',
                'Japanese': 'ja-JP'
            }
            selected_language = st.selectbox("Select audio language:", list(language_options.keys()))
            language = language_options[selected_language]
        
        # Process files when uploaded
        if uploaded_file:
            with st.spinner("Processing file and extracting content..."):
                # Determine file type and extract text
                file_extension = uploaded_file.name.split('.')[-1].lower()
                text = ""
                extracted_images = []
                
                # Create a BytesIO object from uploaded file
                file_stream = io.BytesIO(uploaded_file.read())
                
                if file_extension == 'pdf':
                    text = extract_text_from_pdf(file_stream)
                    # Extract images from PDF
                    file_stream.seek(0)  # Reset stream position
                    extracted_images = extract_images_from_pdf(file_stream)
                    if extracted_images:
                        st.success(f"📸 Extracted {len(extracted_images)} images from PDF!")
                        
                elif file_extension in ['docx', 'doc']:
                    text = extract_text_from_docx(file_stream)
                    # Extract images from DOCX
                    file_stream.seek(0)  # Reset stream position
                    extracted_images = extract_images_from_docx(file_stream)
                    if extracted_images:
                        st.success(f"📸 Extracted {len(extracted_images)} images from DOCX!")
                        
                elif file_extension == 'txt':
                    file_stream.seek(0)
                    text = file_stream.read().decode('utf-8')
                    
                elif file_extension in ['mp3', 'wav', 'm4a', 'ogg', 'flac']:
                    # For audio files, we need to pass the original uploaded file
                    uploaded_file.seek(0)  # Reset file position
                    text = extract_text_from_audio(uploaded_file, language)
                
                if text:
                    # Extract flashcards
                    flashcards = extract_questions_and_answers(text)
                    
                    if flashcards:
                        # Process manually uploaded images
                        manual_images = []
                        if uploaded_images:
                            for img in uploaded_images:
                                img_base64 = process_image_to_base64(img)
                                if img_base64:
                                    manual_images.append(img_base64)
                        
                        # Combine extracted images with manually uploaded images
                        all_images = extracted_images + manual_images
                        
                        # Store in session state
                        st.session_state['flashcards'] = flashcards
                        st.session_state['flashcard_colors'] = [generate_random_color() for _ in flashcards]
                        st.session_state['flashcard_images'] = all_images
                        
                        # Show summary
                        st.success(f"✅ Successfully generated {len(flashcards)} flashcards!")
                        if all_images:
                            st.info(f"🖼️ Total images available: {len(all_images)} (Extracted: {len(extracted_images)}, Manual: {len(manual_images)})")
                        
                        # Show extracted text preview
                        with st.expander("📄 View Extracted Text"):
                            st.text_area("Extracted Content:", text[:1000] + "..." if len(text) > 1000 else text, height=200)
                            
                        # Show extracted images preview
                        if extracted_images:
                            with st.expander("🖼️ View Extracted Images"):
                                cols = st.columns(min(3, len(extracted_images)))
                                for i, img_data in enumerate(extracted_images[:6]):  # Show max 6 images
                                    with cols[i % 3]:
                                        try:
                                            img_bytes = base64.b64decode(img_data)
                                            image = Image.open(io.BytesIO(img_bytes))
                                            st.image(image, caption=f"Extracted Image {i+1}", width=200)
                                        except Exception as e:
                                            st.error(f"Error displaying image {i+1}: {e}")
                                
                                if len(extracted_images) > 6:
                                    st.info(f"+ {len(extracted_images) - 6} more images extracted")
                    else:
                        st.error("❌ Could not extract any flashcards from the uploaded file. Please try a different file or format.")
                else:
                    st.error("❌ Could not extract text from the uploaded file. Please check the file format and try again.")
    
    with tab2:
        # Voice recording section
        create_flashcards_from_voice_notes()
    
    # Display flashcards if available (this section remains the same as before)
    if 'flashcards' in st.session_state and st.session_state['flashcards']:
        st.subheader("🎴 Your Flashcards")
        
        # Navigation controls
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            if st.button("⏮️ Previous") and st.session_state.get('current_card', 0) > 0:
                st.session_state['current_card'] = st.session_state.get('current_card', 0) - 1
                st.rerun()
        
        with col2:
            total_cards = len(st.session_state['flashcards'])
            current_card = st.session_state.get('current_card', 0)
            st.write(f"Card {current_card + 1} of {total_cards}")
            
            # Progress bar
            progress = (current_card + 1) / total_cards
            st.progress(progress)
        
        with col3:
            if st.button("Next ⏭️") and st.session_state.get('current_card', 0) < len(st.session_state['flashcards']) - 1:
                st.session_state['current_card'] = st.session_state.get('current_card', 0) + 1
                st.rerun()
        
        # Display current flashcard
        current_card_index = st.session_state.get('current_card', 0)
        card = st.session_state['flashcards'][current_card_index]
        color = st.session_state['flashcard_colors'][current_card_index]
        
        # Get image for current card (cycle through available images)
        image_data = None
        if st.session_state.get('flashcard_images'):
            image_index = current_card_index % len(st.session_state['flashcard_images'])
            image_data = st.session_state['flashcard_images'][image_index]
        
        display_flashcard(card, color, image_data, current_card_index)
        
        # Study mode controls
        st.subheader("📚 Study Controls")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("🔀 Shuffle Cards"):
                combined = list(zip(st.session_state['flashcards'], st.session_state['flashcard_colors']))
                random.shuffle(combined)
                st.session_state['flashcards'], st.session_state['flashcard_colors'] = zip(*combined)
                st.session_state['flashcards'] = list(st.session_state['flashcards'])
                st.session_state['flashcard_colors'] = list(st.session_state['flashcard_colors'])
                st.session_state['current_card'] = 0
                st.success("Cards shuffled!")
                st.rerun()
        
        with col2:
            if st.button("🎯 Reset to Start"):
                st.session_state['current_card'] = 0
                st.rerun()
        
        with col3:
            if st.button("📊 View All Cards"):
                st.session_state['view_all'] = not st.session_state.get('view_all', False)
                st.rerun()
        
        with col4:
            if st.button("💾 Export JSON"):
                export_data = {
                    'flashcards': st.session_state['flashcards'],
                    'colors': st.session_state['flashcard_colors'],
                    'images': st.session_state.get('flashcard_images', []),
                    'created_at': datetime.now().isoformat()
                }
                json_str = json.dumps(export_data, indent=2)
                st.download_button(
                    label="Download Flashcards",
                    data=json_str,
                    file_name=f"flashcards_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
        
        # View all cards mode
        if st.session_state.get('view_all', False):
            st.subheader("📋 All Flashcards")
            for i, (card, color) in enumerate(zip(st.session_state['flashcards'], st.session_state['flashcard_colors'])):
                with st.expander(f"Card {i+1}: {card['question'][:50]}..."):
                    image_data = None
                    if st.session_state.get('flashcard_images'):
                        image_index = i % len(st.session_state['flashcard_images'])
                        image_data = st.session_state['flashcard_images'][image_index]
                    display_flashcard(card, color, image_data, f"all_{i}")
    
    # Usage instructions
    with st.expander("ℹ️ How to Use"):
        st.markdown("""
        ### Supported Input Methods:
        - **File Upload**: PDF, DOCX, TXT, MP3, WAV, M4A, OGG, FLAC
        - **Voice Recording**: Record your voice directly
        - **Images**: PNG, JPG, JPEG, GIF, BMP (optional reference images)
        
        ### How It Works:
        1. **Upload a file** or **record your voice** with the content you want to study
        2. **Optionally add images** that relate to your study material
        3. **For audio files**, select the appropriate language
        4. The system will automatically extract questions and answers
        5. **Navigate through cards** using the Previous/Next buttons
        6. **Click 'Flip'** to see the answer
        7. Use **study controls** to shuffle, reset, or view all cards
        
        ### Tips for Better Results:
        - For voice recordings, speak clearly and at a moderate pace
        - Structure your content with clear questions and answers
        - Include relevant images to enhance visual memory
        - Use the shuffle feature to test your knowledge randomly
        """)

# def flashcard_generator_page():
#     """Main flashcard generator page."""
#     st.header("🃏 Enhanced Flashcard Generator")
#     st.write("Generate flashcards from text files, documents, audio files, and optionally include images!")
    
#     # File upload section
#     st.subheader("📁 Upload Content")
    
#     # Main content upload
#     uploaded_file = st.file_uploader(
#         "Choose a file (PDF, DOCX, TXT, MP3, WAV, M4A, OGG, FLAC)",
#         type=['pdf', 'docx', 'txt', 'mp3', 'wav', 'm4a', 'ogg', 'flac'],
#         help="Upload text documents or audio files to generate flashcards"
#     )
    
#     # Optional image upload
#     st.subheader("🖼️ Optional: Add Reference Images")
#     uploaded_images = st.file_uploader(
#         "Upload reference images (optional)",
#         type=['png', 'jpg', 'jpeg', 'gif', 'bmp'],
#         accept_multiple_files=True,
#         help="Add images that will be displayed with your flashcards"
#     )
    
#     # Audio language selection (only show if audio file is uploaded)
#     language = 'en-US'
#     if uploaded_file and uploaded_file.name.lower().endswith(('.mp3', '.wav', '.m4a', '.ogg', '.flac')):
#         st.subheader("🗣️ Audio Settings")
#         language_options = {
#             'English (US)': 'en-US',
#             'English (UK)': 'en-GB',
#             'Spanish': 'es-ES',
#             'French': 'fr-FR',
#             'German': 'de-DE',
#             'Italian': 'it-IT',
#             'Portuguese': 'pt-PT',
#             'Russian': 'ru-RU',
#             'Chinese': 'zh-CN',
#             'Japanese': 'ja-JP'
#         }
#         selected_language = st.selectbox("Select audio language:", list(language_options.keys()))
#         language = language_options[selected_language]
    
#     # Process files when uploaded
#     if uploaded_file:
#         with st.spinner("Processing file and extracting content..."):
#             # Determine file type and extract text
#             file_extension = uploaded_file.name.split('.')[-1].lower()
#             text = ""
#             extracted_images = []
            
#             # Create a BytesIO object from uploaded file
#             file_stream = io.BytesIO(uploaded_file.read())
            
#             if file_extension == 'pdf':
#                 text = extract_text_from_pdf(file_stream)
#                 # Extract images from PDF
#                 file_stream.seek(0)  # Reset stream position
#                 extracted_images = extract_images_from_pdf(file_stream)
#                 if extracted_images:
#                     st.success(f"📸 Extracted {len(extracted_images)} images from PDF!")
                    
#             elif file_extension in ['docx', 'doc']:
#                 text = extract_text_from_docx(file_stream)
#                 # Extract images from DOCX
#                 file_stream.seek(0)  # Reset stream position
#                 extracted_images = extract_images_from_docx(file_stream)
#                 if extracted_images:
#                     st.success(f"📸 Extracted {len(extracted_images)} images from DOCX!")
                    
#             elif file_extension == 'txt':
#                 file_stream.seek(0)
#                 text = file_stream.read().decode('utf-8')
                
#             elif file_extension in ['mp3', 'wav', 'm4a', 'ogg', 'flac']:
#                 # For audio files, we need to pass the original uploaded file
#                 uploaded_file.seek(0)  # Reset file position
#                 text = extract_text_from_audio(uploaded_file, language)
            
#             if text:
#                 # Extract flashcards
#                 flashcards = extract_questions_and_answers(text)
                
#                 if flashcards:
#                     # Process manually uploaded images
#                     manual_images = []
#                     if uploaded_images:
#                         for img in uploaded_images:
#                             img_base64 = process_image_to_base64(img)
#                             if img_base64:
#                                 manual_images.append(img_base64)
                    
#                     # Combine extracted images with manually uploaded images
#                     all_images = extracted_images + manual_images
                    
#                     # Store in session state
#                     st.session_state['flashcards'] = flashcards
#                     st.session_state['flashcard_colors'] = [generate_random_color() for _ in flashcards]
#                     st.session_state['flashcard_images'] = all_images
                    
#                     # Show summary
#                     st.success(f"✅ Successfully generated {len(flashcards)} flashcards!")
#                     if all_images:
#                         st.info(f"🖼️ Total images available: {len(all_images)} (Extracted: {len(extracted_images)}, Manual: {len(manual_images)})")
                    
#                     # Show extracted text preview
#                     with st.expander("📄 View Extracted Text"):
#                         st.text_area("Extracted Content:", text[:1000] + "..." if len(text) > 1000 else text, height=200)
                        
#                     # Show extracted images preview
#                     if extracted_images:
#                         with st.expander("🖼️ View Extracted Images"):
#                             cols = st.columns(min(3, len(extracted_images)))
#                             for i, img_data in enumerate(extracted_images[:6]):  # Show max 6 images
#                                 with cols[i % 3]:
#                                     try:
#                                         img_bytes = base64.b64decode(img_data)
#                                         image = Image.open(io.BytesIO(img_bytes))
#                                         st.image(image, caption=f"Extracted Image {i+1}", width=200)
#                                     except Exception as e:
#                                         st.error(f"Error displaying image {i+1}: {e}")
                            
#                             if len(extracted_images) > 6:
#                                 st.info(f"+ {len(extracted_images) - 6} more images extracted")
#                 else:
#                     st.error("❌ Could not extract any flashcards from the uploaded file. Please try a different file or format.")
#             else:
#                 st.error("❌ Could not extract text from the uploaded file. Please check the file format and try again.")
    
#     # Display flashcards if available
#     if 'flashcards' in st.session_state and st.session_state['flashcards']:
#         st.subheader("🎴 Your Flashcards")
        
#         # Navigation controls
#         col1, col2, col3 = st.columns([1, 2, 1])
        
#         with col1:
#             if st.button("⏮️ Previous") and st.session_state.get('current_card', 0) > 0:
#                 st.session_state['current_card'] = st.session_state.get('current_card', 0) - 1
#                 st.rerun()
        
#         with col2:
#             total_cards = len(st.session_state['flashcards'])
#             current_card = st.session_state.get('current_card', 0)
#             st.write(f"Card {current_card + 1} of {total_cards}")
            
#             # Progress bar
#             progress = (current_card + 1) / total_cards
#             st.progress(progress)
        
#         with col3:
#             if st.button("Next ⏭️") and st.session_state.get('current_card', 0) < len(st.session_state['flashcards']) - 1:
#                 st.session_state['current_card'] = st.session_state.get('current_card', 0) + 1
#                 st.rerun()
        
#         # Display current flashcard
#         current_card_index = st.session_state.get('current_card', 0)
#         card = st.session_state['flashcards'][current_card_index]
#         color = st.session_state['flashcard_colors'][current_card_index]
        
#         # Get image for current card (cycle through available images)
#         image_data = None
#         if st.session_state.get('flashcard_images'):
#             image_index = current_card_index % len(st.session_state['flashcard_images'])
#             image_data = st.session_state['flashcard_images'][image_index]
        
#         display_flashcard(card, color, image_data, current_card_index)
        
#         # Study mode controls
#         st.subheader("📚 Study Controls")
#         col1, col2, col3, col4 = st.columns(4)
        
#         with col1:
#             if st.button("🔀 Shuffle Cards"):
#                 combined = list(zip(st.session_state['flashcards'], st.session_state['flashcard_colors']))
#                 random.shuffle(combined)
#                 st.session_state['flashcards'], st.session_state['flashcard_colors'] = zip(*combined)
#                 st.session_state['flashcards'] = list(st.session_state['flashcards'])
#                 st.session_state['flashcard_colors'] = list(st.session_state['flashcard_colors'])
#                 st.session_state['current_card'] = 0
#                 st.success("Cards shuffled!")
#                 st.rerun()
        
#         with col2:
#             if st.button("🎯 Reset to Start"):
#                 st.session_state['current_card'] = 0
#                 st.rerun()
        
#         with col3:
#             if st.button("📊 View All Cards"):
#                 st.session_state['view_all'] = not st.session_state.get('view_all', False)
#                 st.rerun()
        
#         with col4:
#             if st.button("💾 Export JSON"):
#                 export_data = {
#                     'flashcards': st.session_state['flashcards'],
#                     'colors': st.session_state['flashcard_colors'],
#                     'images': st.session_state.get('flashcard_images', []),
#                     'created_at': datetime.now().isoformat()
#                 }
#                 json_str = json.dumps(export_data, indent=2)
#                 st.download_button(
#                     label="Download Flashcards",
#                     data=json_str,
#                     file_name=f"flashcards_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
#                     mime="application/json"
#                 )
        
#         # View all cards mode
#         if st.session_state.get('view_all', False):
#             st.subheader("📋 All Flashcards")
#             for i, (card, color) in enumerate(zip(st.session_state['flashcards'], st.session_state['flashcard_colors'])):
#                 with st.expander(f"Card {i+1}: {card['question'][:50]}..."):
#                     image_data = None
#                     if st.session_state.get('flashcard_images'):
#                         image_index = i % len(st.session_state['flashcard_images'])
#                         image_data = st.session_state['flashcard_images'][image_index]
#                     display_flashcard(card, color, image_data, f"all_{i}")
    
#     # Usage instructions
#     with st.expander("ℹ️ How to Use"):
#         st.markdown("""
#         ### Supported File Types:
#         - **Text Files**: PDF, DOCX, TXT
#         - **Audio Files**: MP3, WAV, M4A, OGG, FLAC
#         - **Images**: PNG, JPG, JPEG, GIF, BMP (optional reference images)
        
#         ### How It Works:
#         1. **Upload a file** containing the content you want to study
#         2. **Optionally add images** that relate to your study material
#         3. **For audio files**, select the appropriate language
#         4. The system will automatically extract questions and answers
#         5. **Navigate through cards** using the Previous/Next buttons
#         6. **Click 'Flip'** to see the answer
#         7. Use **study controls** to shuffle, reset, or view all cards
        
#         ### Tips for Better Results:
#         - Use well-structured documents with clear questions and answers
#         - For audio files, speak clearly and at a moderate pace
#         - Include relevant images to enhance visual memory
#         - Use the shuffle feature to test your knowledge randomly
#         """)

# Initialize session state
if 'current_card' not in st.session_state:
    st.session_state['current_card'] = 0




# === MCQ Generation Functions //////////////////////////////////////////////////////////////////////////////////////===
# 1. Core MCQ Generation
def generate_mcqs_from_text(text, num_questions=10):
    """
    Generate MCQs from text using simple NLP techniques
    This is a basic implementation - you can enhance it with AI/ML models
    """
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    mcqs = []
    used_sentences = set()
    
    for i in range(min(num_questions, len(sentences))):
        # Find sentences that haven't been used
        available_sentences = [s for j, s in enumerate(sentences) if j not in used_sentences and len(s.split()) > 5]
        
        if not available_sentences:
            break
            
        sentence = random.choice(available_sentences)
        used_sentences.add(sentences.index(sentence))
        
        # Simple question generation (this can be enhanced with NLP models)
        words = sentence.split()
        if len(words) < 5:
            continue
            
        # Create a fill-in-the-blank or simple question
        key_word_idx = random.randint(1, len(words) - 2)
        key_word = words[key_word_idx]
        
        # Skip very short words
        if len(key_word) < 3:
            continue
            
        question_text = sentence.replace(key_word, "______")
        
        # Generate wrong options (simple approach)
        wrong_options = generate_wrong_options(key_word, words, text)
        
        options = [key_word] + wrong_options[:3]  # Ensure we have 4 options
        random.shuffle(options)
        
        correct_answer = chr(65 + options.index(key_word))  # A, B, C, D
        
        mcq = {
            'question': f"Fill in the blank: {question_text}",
            'options': {
                'A': options[0],
                'B': options[1],
                'C': options[2],
                'D': options[3] if len(options) > 3 else "None of the above"
            },
            'correct_answer': correct_answer,
            'explanation': f"The correct word from the context is '{key_word}'"
        }
        
        mcqs.append(mcq)
    
    return mcqs

def generate_wrong_options(correct_word, context_words, full_text):
    """Generate plausible wrong options"""
    # Simple approach - you can enhance this
    common_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
    
    # Get words from context
    potential_options = []
    words = full_text.split()
    
    for word in words:
        word = re.sub(r'[^\w\s]', '', word).lower()
        if (len(word) > 2 and 
            word != correct_word.lower() and 
            word not in common_words and
            word not in potential_options):
            potential_options.append(word.capitalize())
    
    # Add some generic wrong options if we don't have enough
    generic_options = ['Therefore', 'However', 'Moreover', 'Furthermore', 'Nevertheless', 
                      'Subsequently', 'Consequently', 'Initially', 'Finally', 'Previously']
    
    potential_options.extend(generic_options)
    
    # Remove duplicates and return 3 random options
    potential_options = list(set(potential_options))
    return random.sample(potential_options, min(3, len(potential_options)))

# === Data Persistence Functions ===
# 2. Data Persistence
def save_mcq_tests_to_file(tests):
    """Save MCQ tests to a pickle file"""
    try:
        with open('mcq_tests.pkl', 'wb') as f:
            pickle.dump(tests, f)
    except Exception as e:
        st.error(f"Error saving tests: {e}")

def load_mcq_tests_from_file():
    """Load MCQ tests from pickle file"""
    try:
        if os.path.exists('mcq_tests.pkl'):
            with open('mcq_tests.pkl', 'rb') as f:
                return pickle.load(f)
    except Exception as e:
        st.error(f"Error loading tests: {e}")
    return []

def save_test_results_to_file(results):
    """Save test results to a pickle file"""
    try:
        with open('test_results.pkl', 'wb') as f:
            pickle.dump(results, f)
    except Exception as e:
        st.error(f"Error saving results: {e}")

def load_test_results_from_file():
    """Load test results from pickle file"""
    try:
        if os.path.exists('test_results.pkl'):
            with open('test_results.pkl', 'rb') as f:
                return pickle.load(f)
    except Exception as e:
        st.error(f"Error loading results: {e}")
    return []

def start_test_with_security(test):
    """Start test with enhanced security measures"""
    st.session_state.current_test = test
    st.session_state.test_mode = 'active'
    st.session_state.test_start_time = time.time()
    st.session_state.current_question = 0
    st.session_state.user_answers = {}
    st.session_state.test_started = False  # Will show fullscreen prompt first
    
    # Reset violation counters
    st.session_state.tab_switches = 0
    st.session_state.copy_attempts = 0
    st.session_state.window_blur = 0
    
    st.rerun()

# 3. Fullscreen Interface   
def inject_fullscreen_script():
    """Inject JavaScript to force fullscreen mode"""
    st.components.v1.html("""
    <script>
    function enterFullscreen() {
        const elem = document.documentElement;
        if (elem.requestFullscreen) {
            elem.requestFullscreen().catch(err => {
                console.error("Fullscreen error:", err);
            });
        } else if (elem.webkitRequestFullscreen) {
            elem.webkitRequestFullscreen();
        } else if (elem.msRequestFullscreen) {
            elem.msRequestFullscreen();
        }
    }
    
    // Try multiple times to ensure fullscreen
    enterFullscreen();
    setTimeout(enterFullscreen, 500);
    setTimeout(enterFullscreen, 1000);
    
    // Also add event listener for when fullscreen might be exited
    document.addEventListener('fullscreenchange', function() {
        if (!document.fullscreenElement) {
            enterFullscreen();
        }
    });
    </script>
    """, height=0)

def inject_fullscreen_and_security_js():
    """Inject JavaScript for fullscreen mode and security measures"""
    js_code = """
    <script>
    // Function to force fullscreen mode
    function enterFullscreen() {
        const elem = document.documentElement;
        if (elem.requestFullscreen) {
            elem.requestFullscreen().catch(err => {
                console.error("Fullscreen error:", err);
                alert("Please allow fullscreen for the test to continue");
            });
        } else if (elem.webkitRequestFullscreen) {
            elem.webkitRequestFullscreen();
        } else if (elem.msRequestFullscreen) {
            elem.msRequestFullscreen();
        } else if (elem.mozRequestFullScreen) {
            elem.mozRequestFullScreen();
        }
    }
    
    // Try to enter fullscreen immediately and repeatedly
    enterFullscreen();
    setTimeout(enterFullscreen, 500);
    setTimeout(enterFullscreen, 1000);
    setTimeout(enterFullscreen, 2000);
    
    // Track security violations
    let violations = {
        tabSwitches: 0,
        copyAttempts: 0,
        fullscreenExits: 0
    };
    
    // Function to show violation warnings
    function showViolation(message) {
        // Create or update warning element
        let warning = document.getElementById('violation-warning');
        if (!warning) {
            warning = document.createElement('div');
            warning.id = 'violation-warning';
            warning.style.cssText = `
                position: fixed;
                top: 20px;
                right: 20px;
                background: #ff4444;
                color: white;
                padding: 10px 15px;
                border-radius: 5px;
                z-index: 9999;
                font-weight: bold;
                box-shadow: 0 2px 10px rgba(0,0,0,0.3);
                animation: fadeIn 0.3s;
            `;
            document.body.appendChild(warning);
        }
        warning.textContent = `⚠️ ${message}`;
        
        // Auto-hide after 3 seconds
        setTimeout(() => {
            warning.style.animation = 'fadeOut 0.3s';
            setTimeout(() => warning.remove(), 300);
        }, 3000);
    }
    
    // Tab switching detection
    document.addEventListener('visibilitychange', function() {
        if (document.hidden) {
            violations.tabSwitches++;
            showViolation(`Tab switch detected (${violations.tabSwitches})`);
            
            // Send to Streamlit
            if (window.parent && window.parent.postMessage) {
                window.parent.postMessage({
                    type: 'violation',
                    violationType: 'tab_switch',
                    count: violations.tabSwitches
                }, '*');
            }
            
            // Force focus back to this window when it becomes visible again
            document.addEventListener('visibilitychange', function onVisibilityChange() {
                if (!document.hidden) {
                    window.focus();
                    enterFullscreen();
                    document.removeEventListener('visibilitychange', onVisibilityChange);
                }
            });
        }
    });
    
    // Fullscreen exit detection
    document.addEventListener('fullscreenchange', function() {
        if (!document.fullscreenElement) {
            violations.fullscreenExits++;
            showViolation(`Fullscreen exit detected (${violations.fullscreenExits})`);
            
            // Try to re-enter fullscreen
            setTimeout(enterFullscreen, 500);
            
            // Send to Streamlit
            if (window.parent && window.parent.postMessage) {
                window.parent.postMessage({
                    type: 'violation',
                    violationType: 'fullscreen_exit',
                    count: violations.fullscreenExits
                }, '*');
            }
        }
    });
    
    // Disable right-click
    document.addEventListener('contextmenu', function(e) {
        e.preventDefault();
        showViolation('Right-click disabled');
        return false;
    });
    
    // Keyboard restrictions
    document.addEventListener('keydown', function(e) {
        // Disable developer tools
        if (e.key === 'F12' || (e.ctrlKey && e.shiftKey && e.key === 'I')) {
            e.preventDefault();
            showViolation('Developer tools disabled');
            return false;
        }
        
        // Disable copy/paste
        if (e.ctrlKey && (e.key === 'c' || e.key === 'C')) {
            e.preventDefault();
            violations.copyAttempts++;
            showViolation(`Copy attempt (${violations.copyAttempts})`);
            return false;
        }
        
        if (e.ctrlKey && (e.key === 'v' || e.key === 'V')) {
            e.preventDefault();
            showViolation('Paste disabled');
            return false;
        }
        
        // Disable text selection
        if (e.ctrlKey && (e.key === 'a' || e.key === 'A')) {
            e.preventDefault();
            showViolation('Text selection disabled');
            return false;
        }
        
        // Disable window switching
        if (e.altKey && e.key === 'Tab') {
            e.preventDefault();
            showViolation('Window switching disabled');
            return false;
        }
        
        // Prevent exiting fullscreen with ESC
        if (e.key === 'Escape') {
            e.preventDefault();
            showViolation('Please remain in fullscreen mode');
            enterFullscreen();
            return false;
        }
    });
    
    // Disable text selection
    document.addEventListener('selectstart', function(e) {
        e.preventDefault();
        return false;
    });
    
    // Disable drag and drop
    document.addEventListener('dragstart', function(e) {
        e.preventDefault();
        return false;
    });
    
    // Add CSS animations
    const style = document.createElement('style');
    style.textContent = `
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(-20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        @keyframes fadeOut {
            from { opacity: 1; transform: translateY(0); }
            to { opacity: 0; transform: translateY(-20px); }
        }
        body {
            -webkit-user-select: none;
            -moz-user-select: none;
            -ms-user-select: none;
            user-select: none;
        }
    `;
    document.head.appendChild(style);
    
    // Handle browser back button
    window.addEventListener('popstate', function(e) {
        e.preventDefault();
        showViolation('Navigation disabled during test');
        history.pushState(null, document.title, window.location.href);
        return false;
    });
    
    // Push initial state to prevent back button
    history.pushState(null, document.title, window.location.href);
    
    // Disable browser refresh
    window.addEventListener('beforeunload', function(e) {
        e.preventDefault();
        e.returnValue = 'Are you sure you want to leave the test? All progress will be lost.';
        return 'Are you sure you want to leave the test? All progress will be lost.';
    });
    
    // Disable browser navigation keys
    window.addEventListener('keydown', function(e) {
        if ((e.key === 'F5') || 
            (e.ctrlKey && e.key === 'r') || 
            (e.ctrlKey && e.key === 'F5')) {
            e.preventDefault();
            showViolation('Page refresh disabled during test');
            return false;
        }
    });
    
    // Disable browser menu
    window.addEventListener('keydown', function(e) {
        if (e.key === 'F10' || 
            (e.shiftKey && e.key === 'F10') || 
            (e.ctrlKey && e.key === 'm')) {
            e.preventDefault();
            showViolation('Browser menu disabled');
            return false;
        }
    });
    
    console.log('Security measures initialized');
    </script>
    """
    components.html(js_code, height=0)

def show_fullscreen_prompt():
    """Show a fullscreen activation prompt before starting the test"""
    st.markdown("""
    <style>
    .fullscreen-prompt {
        background: linear-gradient(135deg, #1e3c72, #2a5298);
        color: white;
        padding: 3rem;
        border-radius: 20px;
        text-align: center;
        margin: 2rem 0;
        box-shadow: 0 15px 35px rgba(0,0,0,0.2);
    }
    .fs-button {
        background: linear-gradient(45deg, #4CAF50, #2E7D32);
        color: white;
        border: none;
        padding: 1rem 2rem;
        font-size: 1.2em;
        border-radius: 50px;
        cursor: pointer;
        transition: all 0.3s;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        margin: 1rem auto;
        display: block;
        width: fit-content;
    }
    .fs-button:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.3);
    }
    </style>
    """, unsafe_allow_html=True)
    
    current_test = st.session_state.current_test
    
    st.markdown(f"""
    <div class="fullscreen-prompt">
        <h1>🔒 Secure Test Environment</h1>
        <h2>{current_test['title']}</h2>
        <p style="font-size: 1.1em;">This test will open in fullscreen mode with anti-cheating measures</p>
        
        <div style="background: rgba(255,255,255,0.1); padding: 1.5rem; border-radius: 10px; margin: 1.5rem 0;">
            <h3>📋 Test Rules:</h3>
            <ul style="text-align: left; max-width: 400px; margin: 0 auto;">
                <li>Stay in fullscreen mode</li>
                <li>No tab switching allowed</li>
                <li>Right-click is disabled</li>
                <li>Copy/paste is disabled</li>
                <li>Violations will be recorded</li>
            </ul>
        </div>
        
        <p><strong>Questions:</strong> {len(current_test['questions'])} | 
           <strong>Estimated Time:</strong> ~{len(current_test['questions'])*2} minutes</p>
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("🚀 START TEST IN FULLSCREEN", key="start_fullscreen_test"):
        st.session_state.test_started = True
        st.rerun()

def show_fullscreen_test_interface():
    """Show the test in fullscreen mode with enhanced security"""
    current_test = st.session_state.current_test
    current_q_idx = st.session_state.get('current_question', 0)
    
    # Inject fullscreen script with enhanced security
    inject_fullscreen_and_security_js()
    
    # Enhanced CSS with better contrast
    st.markdown("""
    <style>
    /* Hide Streamlit elements */
    .stApp > header, .stApp > footer, .stApp > div[data-testid="stToolbar"] {
        display: none !important;
    }
    
    /* Fullscreen styling */
    .main .block-container {
        padding-top: 1rem !important;
        padding-bottom: 1rem !important;
        max-width: 100% !important;
    }
    
    /* Security warning bar */
    .security-bar {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        background: linear-gradient(90deg, #e74c3c, #c0392b);
        color: white;
        text-align: center;
        padding: 0.5rem;
        z-index: 1000;
        font-weight: bold;
        box-shadow: 0 2px 10px rgba(0,0,0,0.2);
    }
    
    /* Question card with better contrast */
    .question-card {
        background: #ffffff !important;
        border-radius: 10px;
        padding: 2rem;
        margin: 1rem 0;
        border-left: 5px solid #3498db;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    /* Make question text black and readable */
    .question-text {
        color: #000000 !important;
        font-size: 1.1em;
        line-height: 1.6;
    }
    
    /* Test container */
    .test-container {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 15px;
        padding: 2rem;
        margin: 1rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.1);
    }
    
    /* Disable text selection */
    * {
        -webkit-user-select: none;
        -moz-user-select: none;
        -ms-user-select: none;
        user-select: none;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Security status bar
    st.markdown(f"""
    <div class="security-bar">
        🔒 SECURE TEST MODE | Fullscreen Required | Question: {current_q_idx + 1}/{len(current_test['questions'])} | 
        Violations: {st.session_state.get('tab_switches', 0)} tab switches, {st.session_state.get('fullscreen_exits', 0)} fullscreen exits
    </div>
    """, unsafe_allow_html=True)
    
    # Add padding for fixed header
    st.markdown("<div style='height: 60px;'></div>", unsafe_allow_html=True)
    
    # Main test content
    with st.container():
        st.markdown(f"""
        <div class="test-container">
            <h1 style="text-align: center; color: #2c3e50;">{current_test['title']}</h1>
        </div>
        """, unsafe_allow_html=True)
        
        # Progress bar
        progress = (current_q_idx + 1) / len(current_test['questions'])
        st.progress(progress)
        
        # Current question with improved contrast
        if current_q_idx < len(current_test['questions']):
            question = current_test['questions'][current_q_idx]
            
            st.markdown(f"""
            <div class="question-card">
                <h2 style="color: #2c3e50;">Question {current_q_idx + 1}</h2>
                <p class="question-text">{question['question']}</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Answer options
            user_answer = st.radio(
                "Select your answer:",
                options=list(question['options'].keys()),
                format_func=lambda x: f"{x}. {question['options'][x]}",
                key=f"q_{current_q_idx}",
                index=None
            )
            
            # Store answer
            if user_answer:
                st.session_state.user_answers[current_q_idx] = user_answer
            
            # Navigation buttons
            col1, col2, col3 = st.columns([1,1,1])
            
            with col1:
                if current_q_idx > 0 and st.button("⬅️ Previous"):
                    st.session_state.current_question = current_q_idx - 1
                    st.rerun()
            
            with col2:
                if st.button("🔄 Clear"):
                    if current_q_idx in st.session_state.user_answers:
                        del st.session_state.user_answers[current_q_idx]
                    st.rerun()
            
            with col3:
                if current_q_idx < len(current_test['questions']) - 1:
                    if st.button("Next ➡️"):
                        st.session_state.current_question = current_q_idx + 1
                        st.rerun()
                else:
                    if st.button("🏁 Finish Test", type="primary"):
                        st.session_state.test_mode = 'review'
                        st.rerun()

# 4. Test Control
def start_test_with_fullscreen(test):
    """Start test with fullscreen initialization"""
    st.session_state.current_test = test
    st.session_state.test_mode = 'active'
    st.session_state.test_start_time = time.time()
    st.session_state.current_question = 0
    st.session_state.user_answers = {}
    st.session_state.test_started = False  # Will show fullscreen prompt first
    
    # Reset violation counters
    st.session_state.tab_switches = 0
    st.session_state.fullscreen_exits = 0
    st.session_state.copy_attempts = 0
    
    st.rerun()

def mcq_test_review():
    """Review answers and show results"""
    
    if 'current_test' not in st.session_state:
        st.error("No test data found!")
        return
    
    current_test = st.session_state.current_test
    user_answers = st.session_state.get('user_answers', {})
    
    # Calculate results
    correct_count = 0
    total_questions = len(current_test['questions'])
    
    for i, question in enumerate(current_test['questions']):
        if i in user_answers and user_answers[i] == question['correct_answer']:
            correct_count += 1
    
    score = correct_count
    percentage = (correct_count / total_questions) * 100
    time_taken = int(time.time() - st.session_state.test_start_time)
    
    # Save results
    result = {
        'test_id': current_test['id'],
        'test_title': current_test['title'],
        'score': score,
        'total_questions': total_questions,
        'correct_answers': correct_count,
        'wrong_answers': total_questions - correct_count,
        'percentage': percentage,
        'time_taken': time_taken,
        'completed_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'answers': user_answers
    }
    
    st.session_state.test_results.append(result)
    save_test_results_to_file(st.session_state.test_results)
    
    # Show results
    st.markdown(f"# 🎯 Test Results: {current_test['title']}")
    
    # Results summary
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📊 Score", f"{score}/{total_questions}")
    with col2:
        st.metric("📈 Percentage", f"{percentage:.1f}%")
    with col3:
        st.metric("✅ Correct", correct_count)
    with col4:
        st.metric("⏱️ Time", f"{time_taken//60}:{time_taken%60:02d}")
    
    # Performance message
    if percentage >= 90:
        st.success("🌟 Excellent! Outstanding performance!")
    elif percentage >= 80:
        st.success("🎉 Great job! Very good performance!")
    elif percentage >= 70:
        st.info("👍 Good work! Room for improvement.")
    elif percentage >= 60:
        st.warning("📚 Fair performance. Consider reviewing the material.")
    else:
        st.error("💪 Keep studying! You can do better next time.")
    
    # Detailed review
    st.markdown("---")
    st.subheader("📋 Detailed Review")
    
    for i, question in enumerate(current_test['questions']):
        user_answer = user_answers.get(i, "Not answered")
        correct_answer = question['correct_answer']
        is_correct = user_answer == correct_answer
        
        # Question container
        status_icon = "✅" if is_correct else "❌" if user_answer != "Not answered" else "⭕"
        
        with st.expander(f"{status_icon} Question {i+1} - {'Correct' if is_correct else 'Wrong' if user_answer != 'Not answered' else 'Not Answered'}"):
            st.write(f"**Question:** {question['question']}")
            
            # Show options with highlighting
            st.write("**Options:**")
            for option, text in question['options'].items():
                if option == correct_answer:
                    st.success(f"✅ {option}. {text} (Correct Answer)")
                elif option == user_answer and user_answer != correct_answer:
                    st.error(f"❌ {option}. {text} (Your Answer)")
                else:
                    st.write(f"   {option}. {text}")
            
            if user_answer == "Not answered":
                st.warning("⭕ You did not answer this question")
            
            if 'explanation' in question:
                st.info(f"💡 **Explanation:** {question['explanation']}")
    
    # Action buttons
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Take Another Test"):
            # Clean up session state
            for key in ['current_test', 'test_mode', 'test_start_time', 'current_question', 'user_answers']:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()
    
    with col2:
        if st.button("📊 View All Results"):
            st.session_state.view_all_results = True
            st.rerun()

# 5. Main Handler & Generator
def mcq_test_handler():
    """Handle the different test modes"""
    if 'current_test' not in st.session_state:
        mcq_test_generator()
        return
    
    test_mode = st.session_state.get('test_mode', None)
    
    if test_mode == 'active':
        if not st.session_state.get('test_started', False):
            show_fullscreen_prompt()
        else:
            show_fullscreen_test_interface()
    elif test_mode == 'review':
        mcq_test_review()
    else:
        mcq_test_generator()

# === MCQ Test Generator Main Page ===
def mcq_test_generator():
    st.header("🧠 MCQ Test Generator")
    
    # Initialize session state
    if 'mcq_tests' not in st.session_state:
        st.session_state.mcq_tests = load_mcq_tests_from_file()
    
    if 'test_results' not in st.session_state:
        st.session_state.test_results = load_test_results_from_file()
    
    # Main test generation section
    st.subheader("📄 Generate New MCQ Test")
    
    # File upload or text input
    uploaded_file = st.file_uploader("Upload PDF/DOCX/TXT file", type=['pdf', 'docx', 'txt'])
    text_input = st.text_area("Or paste your text here", height=200)
    
    col1, col2 = st.columns(2)
    with col1:
        num_questions = st.number_input("Number of Questions", min_value=5, max_value=50, value=10)
    with col2:
        test_title = st.text_input("Test Title", value=f"Test {len(st.session_state.mcq_tests) + 1}")
    
    if st.button("🎯 Generate MCQ Test", type="primary"):
        text_content = ""
        
        # Extract text from uploaded file
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                # You'll need to implement this function or use a library like PyPDF2
                st.error("PDF extraction not implemented. Please paste text directly.")
            elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                # You'll need to implement this function or use a library like python-docx
                st.error("DOCX extraction not implemented. Please paste text directly.")
            elif uploaded_file.type == "text/plain":
                text_content = uploaded_file.getvalue().decode("utf-8")
        elif text_input.strip():
            text_content = text_input
        
        if not text_content.strip():
            st.error("Please provide text content or upload a file!")
            return
        
        # Generate MCQs
        with st.spinner("Generating MCQ test... 🤔"):
            mcqs = generate_mcqs_from_text(text_content, num_questions)
            
            if mcqs:
                # Create test object
                new_test = {
                    'id': len(st.session_state.mcq_tests) + 1,
                    'title': test_title,
                    'questions': mcqs,
                    'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'status': 'generated',  # generated, scheduled, completed
                    'scheduled_date': None,
                    'scheduled_time': None
                }
                
                st.session_state.mcq_tests.append(new_test)
                save_mcq_tests_to_file(st.session_state.mcq_tests)
                
                st.success("🎉 Test Generated Successfully!!")
                
                # Action buttons
                st.markdown("### Choose Your Next Action:")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("🚀 Start Now", type="primary"):
                        start_test_with_fullscreen(new_test)
                
                with col2:
                    if st.button("📅 Schedule Later"):
                        st.session_state.scheduling_test = new_test['id']
                        st.rerun()
                
                with col3:
                    if st.button("🤔 Decide Later"):
                        st.info("Test saved! You can start it anytime from the dashboard below.")
            else:
                st.error("Failed to generate MCQs. Please try with different content.")
    
    # Schedule test form
    if st.session_state.get('scheduling_test'):
        test_to_schedule = next((test for test in st.session_state.mcq_tests if test['id'] == st.session_state.scheduling_test), None)
        
        if test_to_schedule:
            st.markdown("---")
            st.subheader(f"📅 Schedule Test: {test_to_schedule['title']}")
            
            col1, col2 = st.columns(2)
            with col1:
                scheduled_date = st.date_input("Test Date", min_value=date.today())
            with col2:
                scheduled_time = st.time_input("Test Time")
            
            col_save, col_cancel = st.columns(2)
            with col_save:
                if st.button("💾 Save Schedule"):
                    # Update test with schedule
                    for i, test in enumerate(st.session_state.mcq_tests):
                        if test['id'] == st.session_state.scheduling_test:
                            st.session_state.mcq_tests[i]['status'] = 'scheduled'
                            st.session_state.mcq_tests[i]['scheduled_date'] = scheduled_date.strftime("%Y-%m-%d")
                            st.session_state.mcq_tests[i]['scheduled_time'] = scheduled_time.strftime("%H:%M")
                            break
                    
                    save_mcq_tests_to_file(st.session_state.mcq_tests)
                    st.session_state.scheduling_test = None
                    st.success("Test scheduled successfully!")
                    st.rerun()
            
            with col_cancel:
                if st.button("❌ Cancel"):
                    st.session_state.scheduling_test = None
                    st.rerun()
    
    # Dashboard
    st.markdown("---")
    st.subheader("📊 Test Dashboard")
    
    # Statistics
    total_tests = len([test for test in st.session_state.mcq_tests])
    completed_tests = len([result for result in st.session_state.test_results])
    scheduled_tests = len([test for test in st.session_state.mcq_tests if test['status'] == 'scheduled'])
    pending_tests = len([test for test in st.session_state.mcq_tests if test['status'] == 'generated'])
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📝 Total Tests", total_tests)
    with col2:
        st.metric("✅ Completed", completed_tests)
    with col3:
        st.metric("📅 Scheduled", scheduled_tests)
    with col4:
        st.metric("⏳ Pending", pending_tests)
    
    # Scheduled Tests
    scheduled_tests_list = [test for test in st.session_state.mcq_tests if test['status'] == 'scheduled']
    if scheduled_tests_list:
        st.markdown("### 📅 Scheduled Tests")
        for test in scheduled_tests_list:
            with st.expander(f"{test['title']} - {test['scheduled_date']} at {test['scheduled_time']}"):
                col1, col2, col3 = st.columns([2, 1, 1])
                with col1:
                    st.write(f"**Questions:** {len(test['questions'])}")
                    st.write(f"**Created:** {test['created_at']}")
                with col2:
                    if st.button("🚀 Start Now", key=f"start_scheduled_{test['id']}"):
                        start_test_with_fullscreen(test)
                with col3:
                    if st.button("🗑️ Delete", key=f"delete_scheduled_{test['id']}"):
                        st.session_state.mcq_tests = [t for t in st.session_state.mcq_tests if t['id'] != test['id']]
                        save_mcq_tests_to_file(st.session_state.mcq_tests)
                        st.rerun()
    
    # Pending Tests (Decide Later)
    pending_tests_list = [test for test in st.session_state.mcq_tests if test['status'] == 'generated']
    if pending_tests_list:
        st.markdown("### ⏳ Tests Awaiting Decision")
        for test in pending_tests_list:
            with st.expander(f"{test['title']} - Created {test['created_at']}"):
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button("🚀 Start Now", key=f"start_pending_{test['id']}"):
                        start_test_with_fullscreen(test)
                with col2:
                    if st.button("📅 Schedule", key=f"schedule_pending_{test['id']}"):
                        st.session_state.scheduling_test = test['id']
                        st.rerun()
                with col3:
                    if st.button("👁️ Preview", key=f"preview_{test['id']}"):
                        st.session_state.preview_test = test['id']
                        st.rerun()
                with col4:
                    if st.button("🗑️ Delete", key=f"delete_pending_{test['id']}"):
                        st.session_state.mcq_tests = [t for t in st.session_state.mcq_tests if t['id'] != test['id']]
                        save_mcq_tests_to_file(st.session_state.mcq_tests)
                        st.rerun()
    
    # Test Results History
    if st.session_state.test_results:
        st.markdown("### 📈 Test Results History")
        for result in reversed(st.session_state.test_results[-5:]):  # Show last 5 results
            with st.expander(f"{result['test_title']} - Score: {result['score']}/{result['total_questions']} ({result['percentage']:.1f}%)"):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Date:** {result['completed_at']}")
                    st.write(f"**Time Taken:** {result['time_taken']} seconds")
                with col2:
                    st.write(f"**Correct:** {result['correct_answers']}")
                    st.write(f"**Wrong:** {result['wrong_answers']}")
    
    # Preview test modal
    if st.session_state.get('preview_test'):
        test_to_preview = next((test for test in st.session_state.mcq_tests if test['id'] == st.session_state.preview_test), None)
        if test_to_preview:
            st.markdown("---")
            st.subheader(f"👁️ Preview: {test_to_preview['title']}")
            
            for i, q in enumerate(test_to_preview['questions'][:3]):  # Show first 3 questions
                st.write(f"**Question {i+1}:** {q['question']}")
                for option, text in q['options'].items():
                    st.write(f"   {option}. {text}")
                st.write(f"   **Correct Answer:** {q['correct_answer']}")
                st.write("---")
            
            if len(test_to_preview['questions']) > 3:
                st.write(f"... and {len(test_to_preview['questions']) - 3} more questions")
            
            if st.button("❌ Close Preview"):
                st.session_state.preview_test = None
                st.rerun()

# === Full Screen MCQ Test Interface ===
# def mcq_test_interface():
#     """Full screen MCQ test interface with anti-cheating measures"""
    
#     if 'current_test' not in st.session_state or 'test_mode' not in st.session_state:
#         st.error("No active test found!")
#         return
    
#     current_test = st.session_state.current_test
#     current_q_idx = st.session_state.get('current_question', 0)
    
#     # Initialize cheating detection counters
#     if 'tab_switches' not in st.session_state:
#         st.session_state.tab_switches = 0
#     if 'fullscreen_exits' not in st.session_state:
#         st.session_state.fullscreen_exits = 0
#     if 'violations' not in st.session_state:
#         st.session_state.violations = {'tab_switches': 0, 'fullscreen_exits': 0, 'copy_attempts': 0}
    
#     # Check if this is the first load of the test
#     if 'test_initialized' not in st.session_state:
#         st.session_state.test_initialized = True
#         show_fullscreen_prompt()
#         return

#     # Enhanced anti-cheating CSS and JavaScript
#     st.markdown("""
#     <style>
#     /* Hide Streamlit elements during test */
#     .main > div {
#         padding-top: 0rem;
#     }
#     .stApp > header {
#         display: none !important;
#     }
#     .stApp > footer {
#         display: none !important;
#     }
    
#     /* Test styling */
#     .stApp {
#         background-color: #1e1e1e;
#         color: white;
#     }
#     .test-container {
#         background-color: #2d2d2d;
#         padding: 2rem;
#         border-radius: 10px;
#         margin: 1rem 0;
#         box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
#     }
#     .question-header {
#         background-color: #4a4a4a;
#         padding: 1.5rem;
#         border-radius: 8px;
#         margin-bottom: 1rem;
#         border-left: 4px solid #2196f3;
#     }
#     .test-warning {
#         position: fixed;
#         top: 0;
#         left: 0;
#         width: 100%;
#         background: linear-gradient(90deg, #f44336, #e91e63);
#         color: white;
#         text-align: center;
#         padding: 1rem;
#         z-index: 1000;
#         font-weight: bold;
#         font-size: 1.1em;
#         box-shadow: 0 2px 10px rgba(0,0,0,0.3);
#     }
#     .violation-warning {
#         position: fixed;
#         top: 80px;
#         right: 20px;
#         background-color: #ff5722;
#         color: white;
#         padding: 1rem;
#         border-radius: 8px;
#         z-index: 999;
#         font-weight: bold;
#         animation: pulse 2s infinite;
#     }
#     @keyframes pulse {
#         0% { opacity: 1; }
#         50% { opacity: 0.7; }
#         100% { opacity: 1; }
#     }
    
#     /* Disable text selection */
#     body {
#         -webkit-user-select: none;
#         -moz-user-select: none;
#         -ms-user-select: none;
#         user-select: none;
#     }
    
#     /* Hide scrollbars */
#     ::-webkit-scrollbar {
#         width: 8px;
#     }
#     ::-webkit-scrollbar-track {
#         background: #2d2d2d;
#     }
#     ::-webkit-scrollbar-thumb {
#         background: #555;
#         border-radius: 4px;
#     }
#     </style>
#     """, unsafe_allow_html=True)
    
#     # JavaScript for anti-cheating measures
#     st.components.v1.html(f"""
#     <script>
#     console.log("Initializing test security measures...");
    
#     // Variables to track violations
#     let tabSwitches = {st.session_state.tab_switches};
#     let fullscreenExits = {st.session_state.fullscreen_exits};
#     let warningShown = false;
    
#     // Function to enter fullscreen
#     function enterFullscreen() {{
#         try {{
#             const elem = document.documentElement;
#             if (elem.requestFullscreen) {{
#                 elem.requestFullscreen();
#             }} else if (elem.webkitRequestFullscreen) {{
#                 elem.webkitRequestFullscreen();
#             }} else if (elem.msRequestFullscreen) {{
#                 elem.msRequestFullscreen();
#             }} else if (elem.mozRequestFullScreen) {{
#                 elem.mozRequestFullScreen();
#             }}
#             console.log("Fullscreen requested");
#         }} catch (error) {{
#             console.error("Fullscreen error:", error);
#             alert("⚠️ Please manually enter fullscreen mode (F11) for the test.");
#         }}
#     }}
    
#     // Function to show violation warning
#     function showViolationWarning(message) {{
#         // Remove existing warning
#         const existingWarning = document.querySelector('.violation-warning');
#         if (existingWarning) {{
#             existingWarning.remove();
#         }}
        
#         // Create new warning
#         const warning = document.createElement('div');
#         warning.className = 'violation-warning';
#         warning.innerHTML = message;
#         document.body.appendChild(warning);
        
#         // Remove after 5 seconds
#         setTimeout(() => {{
#             warning.remove();
#         }}, 5000);
#     }}
    
#     // Tab switching detection
#     document.addEventListener('visibilitychange', function() {{
#         if (document.hidden) {{
#             tabSwitches++;
#             console.log("Tab switch detected:", tabSwitches);
#             showViolationWarning(`⚠️ Tab switching detected! Count: ${{tabSwitches}}`);
            
#             // Store in Streamlit session (this is a workaround since we can't directly update session state)
#             if (window.parent && window.parent.postMessage) {{
#                 window.parent.postMessage({{
#                     type: 'tab_switch',
#                     count: tabSwitches
#                 }}, '*');
#             }}
#         }}
#     }});
    
#     // Fullscreen change detection
#     document.addEventListener('fullscreenchange', function() {{
#         if (!document.fullscreenElement) {{
#             fullscreenExits++;
#             console.log("Fullscreen exit detected:", fullscreenExits);
#             showViolationWarning(`⚠️ Fullscreen exit detected! Count: ${{fullscreenExits}}`);
            
#             // Try to re-enter fullscreen after a short delay
#             setTimeout(() => {{
#                 if (!document.fullscreenElement) {{
#                     enterFullscreen();
#                 }}
#             }}, 2000);
#         }}
#     }});
    
#     // Disable right-click
#     document.addEventListener('contextmenu', function(e) {{
#         e.preventDefault();
#         showViolationWarning('⚠️ Right-click is disabled during the test!');
#         return false;
#     }});
    
#     // Disable common keyboard shortcuts
#     document.addEventListener('keydown', function(e) {{
#         // F12 - Developer Tools
#         if (e.key === 'F12') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Developer tools access is restricted!');
#             return false;
#         }}
        
#         // Ctrl+Shift+I - Developer Tools
#         if (e.ctrlKey && e.shiftKey && e.key === 'I') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Developer tools access is restricted!');
#             return false;
#         }}
        
#         // Ctrl+Shift+C - Inspect Element
#         if (e.ctrlKey && e.shiftKey && e.key === 'C') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Inspect element is disabled!');
#             return false;
#         }}
        
#         // Ctrl+U - View Source
#         if (e.ctrlKey && e.key === 'u') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ View source is disabled!');
#             return false;
#         }}
        
#         // Ctrl+A - Select All
#         if (e.ctrlKey && e.key === 'a') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Text selection is disabled!');
#             return false;
#         }}
        
#         // Ctrl+C - Copy
#         if (e.ctrlKey && e.key === 'c') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Copying is disabled during the test!');
#             return false;
#         }}
        
#         // Ctrl+V - Paste
#         if (e.ctrlKey && e.key === 'v') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Pasting is disabled during the test!');
#             return false;
#         }}
        
#         // Alt+Tab - Window switching
#         if (e.altKey && e.key === 'Tab') {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Window switching is not allowed!');
#             return false;
#         }}
        
#         // F11 - Fullscreen toggle (allow this one)
#         // Escape key handling
#         if (e.key === 'Escape' && document.fullscreenElement) {{
#             e.preventDefault();
#             showViolationWarning('⚠️ Please remain in fullscreen mode!');
#             return false;
#         }}
#     }});
    
#     // Disable text selection with mouse
#     document.addEventListener('selectstart', function(e) {{
#         e.preventDefault();
#         return false;
#     }});
    
#     // Disable drag and drop
#     document.addEventListener('dragstart', function(e) {{
#         e.preventDefault();
#         return false;
#     }});
    
#     // Print detection
#     window.addEventListener('beforeprint', function(e) {{
#         e.preventDefault();
#         showViolationWarning('⚠️ Printing is disabled during the test!');
#         return false;
#     }});
    
#     // Focus detection
#     window.addEventListener('blur', function() {{
#         console.log("Window lost focus");
#         showViolationWarning('⚠️ Please keep the test window focused!');
#     }});
    
#     // Auto-enter fullscreen when the page loads
#     window.addEventListener('load', function() {{
#         console.log("Page loaded, entering fullscreen...");
#         setTimeout(() => {{
#             enterFullscreen();
#         }}, 1000);
#     }});
    
#     // Enter fullscreen immediately
#     setTimeout(() => {{
#         console.log("Attempting immediate fullscreen...");
#         enterFullscreen();
#     }}, 500);
    
#     console.log("Test security measures initialized");
#     </script>
#     """, height=0)
    
#     # Test header with warnings
#     st.markdown(f"""
#     <div class="test-warning">
#         🔒 SECURE TEST MODE ACTIVE - Fullscreen Required | Tab Switches: {st.session_state.tab_switches} | Fullscreen Exits: {st.session_state.fullscreen_exits}
#     </div>
#     """, unsafe_allow_html=True)
    
#     # Add some space for the fixed header
#     st.markdown("<br><br>", unsafe_allow_html=True)
    
#     # Progress bar
#     progress = (current_q_idx + 1) / len(current_test['questions'])
#     st.progress(progress)
    
#     # Test info header
#     col1, col2, col3 = st.columns([2, 1, 1])
#     with col1:
#         st.markdown(f"## 📝 {current_test['title']}")
#     with col2:
#         st.markdown(f"**Question:** {current_q_idx + 1}/{len(current_test['questions'])}")
#     with col3:
#         elapsed = int(time.time() - st.session_state.test_start_time)
#         st.markdown(f"**Time:** {elapsed//60}:{elapsed%60:02d}")
    
#     # Show violations if any
#     if st.session_state.tab_switches > 0 or st.session_state.fullscreen_exits > 0:
#         st.warning(f"⚠️ Security Violations Detected: {st.session_state.tab_switches} tab switches, {st.session_state.fullscreen_exits} fullscreen exits")
    
#     # Current question
#     if current_q_idx < len(current_test['questions']):
#         question = current_test['questions'][current_q_idx]
        
#         st.markdown(f"""
#         <div class="test-container">
#             <div class="question-header">
#                 <h3>Question {current_q_idx + 1}</h3>
#                 <p style="font-size: 1.2em; margin-bottom: 0;">{question['question']}</p>
#             </div>
#         </div>
#         """, unsafe_allow_html=True)
        
#         # Options with better styling
#         user_answer = st.radio(
#             "Choose your answer:",
#             options=list(question['options'].keys()),
#             format_func=lambda x: f"{x}. {question['options'][x]}",
#             key=f"q_{current_q_idx}",
#             index=None
#         )
        
#         # Store answer
#         if user_answer:
#             st.session_state.user_answers[current_q_idx] = user_answer
        
#         # Navigation buttons
#         col1, col2, col3 = st.columns([1, 1, 1])
        
#         with col1:
#             if current_q_idx > 0:
#                 if st.button("⬅️ Previous", key="prev_btn"):
#                     st.session_state.current_question = current_q_idx - 1
#                     st.rerun()
        
#         with col2:
#             if st.button("🔄 Clear Answer", key="clear_btn"):
#                 if current_q_idx in st.session_state.user_answers:
#                     del st.session_state.user_answers[current_q_idx]
#                 st.rerun()
        
#         with col3:
#             if current_q_idx < len(current_test['questions']) - 1:
#                 if st.button("Next ➡️", key="next_btn"):
#                     st.session_state.current_question = current_q_idx + 1
#                     st.rerun()
#             else:
#                 if st.button("🏁 Finish Test", type="primary", key="finish_btn"):
#                     st.session_state.test_mode = 'review'
#                     st.rerun()
    
#     # Question navigator
#     st.markdown("### 📋 Question Navigator")
    
#     # Create rows of 10 questions each
#     questions_per_row = 10
#     total_questions = len(current_test['questions'])
    
#     for row_start in range(0, total_questions, questions_per_row):
#         row_end = min(row_start + questions_per_row, total_questions)
#         cols = st.columns(row_end - row_start)
        
#         for i in range(row_start, row_end):
#             with cols[i - row_start]:
#                 status = "✅" if i in st.session_state.user_answers else "⭕"
#                 button_text = f"{status} {i+1}"
                
#                 if i == current_q_idx:
#                     st.markdown(f"**🔸 {button_text}**")
#                 else:
#                     if st.button(button_text, key=f"nav_{i}"):
#                         st.session_state.current_question = i
#                         st.rerun()
    
#     # Summary
#     answered = len(st.session_state.user_answers)
#     unanswered = total_questions - answered
    
#     st.markdown(f"""
#     ### 📊 Progress Summary
#     - **Answered:** {answered}/{total_questions}
#     - **Remaining:** {unanswered}
#     - **Progress:** {(answered/total_questions)*100:.1f}%
#     """)
    
#     # Emergency exit (with strong warning)
#     st.markdown("---")
#     st.markdown("### ⚠️ Emergency Exit")
    
#     if st.checkbox("I understand that exiting will lose all progress"):
#         if st.button("🚪 EXIT TEST (⚠️ ALL PROGRESS WILL BE LOST)", type="secondary"):
#             # Clean up session state
#             for key in ['current_test', 'test_mode', 'test_start_time', 'current_question', 'user_answers', 'tab_switches', 'fullscreen_exits']:
#                 if key in st.session_state:
#                     del st.session_state[key]
#             st.rerun()
#     # Main test interface
#     render_test_interface(current_test, current_q_idx)

#     # if st.button("🚀 Start Test"):
#     #     test = {
#     #         'title': 'Sample Test',
#     #         'questions': [
#     #             {
#     #                 'question': 'What is 2+2?',
#     #                 'options': {'A': '3', 'B': '4', 'C': '5', 'D': '6'},
#     #                 'correct_answer': 'B'
#     #             }
#     #             # Add more questions...
#     #         ]
#     #     }
#     #     start_test_with_security(test)

# def render_test_interface(current_test, current_q_idx):
#     """Render the actual test interface"""
    
#     # Anti-cheating JavaScript that runs during the test
#     anti_cheat_js = """
#     <script>
#         console.log('Initializing test security...');
        
#         let violations = {
#             tab_switches: 0,
#             copy_attempts: 0,
#             fullscreen_exits: 0
#         };
        
#         // Tab visibility detection
#         document.addEventListener('visibilitychange', function() {
#             if (document.hidden) {
#                 violations.tab_switches++;
#                 showViolation('Tab Switch Detected! Count: ' + violations.tab_switches);
                
#                 // Send to Streamlit
#                 window.parent.postMessage({
#                     type: 'violation',
#                     violation_type: 'tab_switch',
#                     count: violations.tab_switches
#                 }, '*');
#             }
#         });
        
#         // Disable right-click
#         document.addEventListener('contextmenu', function(e) {
#             e.preventDefault();
#             showViolation('Right-click is disabled!');
#             return false;
#         });
        
#         // Keyboard restrictions
#         document.addEventListener('keydown', function(e) {
#             // F12, Ctrl+Shift+I (DevTools)
#             if (e.key === 'F12' || (e.ctrlKey && e.shiftKey && e.key === 'I')) {
#                 e.preventDefault();
#                 showViolation('Developer tools are blocked!');
#                 return false;
#             }
            
#             // Ctrl+C (Copy)
#             if (e.ctrlKey && e.key === 'c') {
#                 e.preventDefault();
#                 violations.copy_attempts++;
#                 showViolation('Copying is disabled! Count: ' + violations.copy_attempts);
#                 return false;
#             }
            
#             // Ctrl+V (Paste)
#             if (e.ctrlKey && e.key === 'v') {
#                 e.preventDefault();
#                 showViolation('Pasting is disabled!');
#                 return false;
#             }
            
#             // Ctrl+A (Select All)
#             if (e.ctrlKey && e.key === 'a') {
#                 e.preventDefault();
#                 showViolation('Text selection is disabled!');
#                 return false;
#             }
            
#             // Alt+Tab
#             if (e.altKey && e.key === 'Tab') {
#                 e.preventDefault();
#                 showViolation('Window switching is blocked!');
#                 return false;
#             }
#         });
        
#         // Disable text selection
#         document.addEventListener('selectstart', function(e) {
#             e.preventDefault();
#             return false;
#         });
        
#         // Show violation popup
#         function showViolation(message) {
#             // Remove existing popup
#             const existing = document.querySelector('.violation-popup');
#             if (existing) existing.remove();
            
#             // Create popup
#             const popup = document.createElement('div');
#             popup.className = 'violation-popup';
#             popup.style.cssText = `
#                 position: fixed;
#                 top: 20px;
#                 right: 20px;
#                 background: #f44336;
#                 color: white;
#                 padding: 1rem;
#                 border-radius: 8px;
#                 z-index: 9999;
#                 font-weight: bold;
#                 animation: slideIn 0.3s ease;
#             `;
#             popup.textContent = '⚠️ ' + message;
#             document.body.appendChild(popup);
            
#             // Remove after 3 seconds
#             setTimeout(() => popup.remove(), 3000);
#         }
        
#         // CSS for animations
#         const style = document.createElement('style');
#         style.textContent = `
#             @keyframes slideIn {
#                 from { transform: translateX(100%); opacity: 0; }
#                 to { transform: translateX(0); opacity: 1; }
#             }
#             body { -webkit-user-select: none; -moz-user-select: none; -ms-user-select: none; user-select: none; }
#         `;
#         document.head.appendChild(style);
        
#         console.log('Test security initialized');
#     </script>
#     """
    
#     # Inject the anti-cheating JavaScript
#     components.html(anti_cheat_js, height=0)
    
#     # Test interface styling
#     st.markdown("""
#     <style>
#     .stApp > header { display: none !important; }
#     .stApp > footer { display: none !important; }
#     .main > div { padding-top: 1rem; }
    
#     .test-header {
#         background: linear-gradient(90deg, #1e3c72, #2a5298);
#         color: white;
#         padding: 1rem;
#         border-radius: 10px;
#         margin-bottom: 1rem;
#         text-align: center;
#     }
    
#     .question-container {
#         background: #f8f9fa;
#         padding: 2rem;
#         border-radius: 15px;
#         border-left: 5px solid #2196f3;
#         margin: 1rem 0;
#         box-shadow: 0 4px 6px rgba(0,0,0,0.1);
#     }
    
#     .violation-display {
#         background: #fff3cd;
#         border: 1px solid #ffeaa7;
#         color: #856404;
#         padding: 0.75rem;
#         border-radius: 8px;
#         margin: 1rem 0;
#     }
#     </style>
#     """, unsafe_allow_html=True)
    
#     # Test header
#     st.markdown(f"""
#     <div class="test-header">
#         <h2>🔒 {current_test['title']} - SECURE MODE</h2>
#         <p>Question {current_q_idx + 1} of {len(current_test['questions'])} | 
#         Violations: {st.session_state.violations['tab_switches']} tab switches, 
#         {st.session_state.violations['copy_attempts']} copy attempts</p>
#     </div>
#     """, unsafe_allow_html=True)
    
#     # Progress bar
#     progress = (current_q_idx + 1) / len(current_test['questions'])
#     st.progress(progress)
    
#     # Show violations if any
#     total_violations = sum(st.session_state.violations.values())
#     if total_violations > 0:
#         st.markdown(f"""
#         <div class="violation-display">
#             ⚠️ <strong>Security Violations Detected:</strong> 
#             {st.session_state.violations['tab_switches']} tab switches, 
#             {st.session_state.violations['copy_attempts']} copy attempts, 
#             {st.session_state.violations['fullscreen_exits']} fullscreen exits
#         </div>
#         """, unsafe_allow_html=True)
    
#     # Current question
#     if current_q_idx < len(current_test['questions']):
#         question = current_test['questions'][current_q_idx]
        
#         st.markdown(f"""
#         <div class="question-container">
#             <h3>Question {current_q_idx + 1}</h3>
#             <p style="font-size: 1.1em;">{question['question']}</p>
#         </div>
#         """, unsafe_allow_html=True)
        
#         # Answer options
#         user_answer = st.radio(
#             "Select your answer:",
#             options=list(question['options'].keys()),
#             format_func=lambda x: f"{x}. {question['options'][x]}",
#             key=f"q_{current_q_idx}",
#             index=None
#         )
        
#         # Store answer
#         if user_answer:
#             st.session_state.user_answers[current_q_idx] = user_answer
        
#         # Navigation
#         col1, col2, col3 = st.columns([1,1,1])
        
#         with col1:
#             if current_q_idx > 0 and st.button("⬅️ Previous"):
#                 st.session_state.current_question = current_q_idx - 1
#                 st.rerun()
        
#         with col2:
#             if st.button("🔄 Clear"):
#                 if current_q_idx in st.session_state.user_answers:
#                     del st.session_state.user_answers[current_q_idx]
#                 st.rerun()
        
#         with col3:
#             if current_q_idx < len(current_test['questions']) - 1:
#                 if st.button("Next ➡️"):
#                     st.session_state.current_question = current_q_idx + 1
#                     st.rerun()
#             else:
#                 if st.button("🏁 Finish Test", type="primary"):
#                     st.session_state.test_mode = 'review'
#                     st.rerun()
        
#         # Question navigator
#         st.markdown("### Question Navigator")
#         answered = len(st.session_state.user_answers)
#         st.write(f"Progress: {answered}/{len(current_test['questions'])} questions answered")
        
#         # Emergency exit
#         st.markdown("---")
#         if st.checkbox("Emergency Exit (loses progress)"):
#             if st.button("EXIT TEST"):
#                 for key in ['current_test', 'test_mode', 'test_start_time', 'current_question', 'user_answers', 'violations', 'test_initialized']:
#                     if key in st.session_state:
#                         del st.session_state[key]
#                 st.rerun()

# def start_test_with_fullscreen(test):
    """Start test with fullscreen initialization and security"""
    st.session_state.current_test = test
    st.session_state.test_mode = 'active'
    st.session_state.test_start_time = time.time()
    st.session_state.current_question = 0
    st.session_state.user_answers = {}
    
    # Initialize security counters
    st.session_state.tab_switches = 0
    st.session_state.fullscreen_exits = 0
    st.session_state.copy_attempts = 0
    
    # Force showing the fullscreen interface immediately
    st.session_state.test_started = True
    st.rerun()

def update_main_navigation():
    """Add MCQ Test Generator to the main navigation"""





# === Main //////////////////////////////////////////////////////////////////////////////////////===
def main():
    st.sidebar.title("Study Buddy 📚")
    current_hour = datetime.now().hour
    greeting = "Good Morning" if current_hour < 12 else "Good Afternoon" if current_hour < 17 else "Good Evening"
    st.sidebar.markdown(f"### {greeting}! 👋")

    pages = {
        "📅 Home": calendar_home,
        "📝 Text Summarizer": text_summarizer_page,
        "🃏 Flashcard Generator": flashcard_generator_page,
        "🧠 MCQ Test Generator": mcq_test_handler,
        "🔗 URL Shortener": url_shortener_page,
        "📊 Grade Calculator": grade_calculator,
        "🎓 SGPA Calculator": sgpa_calculator,
        "CGPA Calculator": cgpa_calculator,
        "⚡ Physics Solver": physics_solver,
        "📏 Unit Converter": unit_converter
    }

    choice = st.sidebar.radio("Let's get productive! 🚀", list(pages.keys()))
    
    # Add Settings at the bottom
    st.sidebar.markdown("---")
    if st.sidebar.button("⚙️ Settings"):
        st.session_state.current_page = "Settings"
    
    # Handle page navigation
    if st.session_state.get('current_page') == "Settings":
        voice_settings_page()
    else:
        # Handle voice commands
        handle_voice_commands()
        
        # Add voice indicator to sidebar
        add_voice_to_sidebar()
        
        # Run the selected page
        pages[choice]()

if __name__ == "__main__":
    main()